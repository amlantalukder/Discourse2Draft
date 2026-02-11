from shiny import reactive
from shiny.express import render, module, ui
from shiny.types import ImgData
from utils import Config, print_func_name
from pathlib import Path
import pandas as pd
from .defaults import SpecialSectionTypes, ContentTypes
from ..backend.db import selectFromDB, updateDB, insertIntoDB, \
                vector_db_collections_status, \
                vector_db_collections_type, \
                generated_files_status, \
                generated_files_ai_architecture, \
                uploaded_files_status
from ..backend.vectordb import ChromaDB, deleteCollection, getLoader
from ..backend.ai.tools.search_pubmed import formatAPA
from datetime import datetime
import json
import re
from docx import Document
import textwrap
import logging

@module
def confirmBox(input, output, session, text, action_btn_name, cancel_btn_name):

    agree = False

    with ui.hold() as content:
        with ui.div():
            text
        with ui.div():
            ui.input_action_button(id='btn_action', label=action_btn_name)
            ui.input_action_button(id='btn_cancel', label=cancel_btn_name)
    
    @reactive.effect
    @reactive.event(input.action_btn_name)
    def action():
        agree = True

    return content, agree

@print_func_name
def initProfile(config_app):
    
    if config_app.email != '':
        records = selectFromDB('settings', 
                    field_names=['email'],
                    field_values=[[config_app.email]],
                    order_by_field_names=['update_date'],
                    order_by_types=['DESC'],
                    limit=1)
    else:
        records = selectFromDB('settings', 
                    field_names=['session'],
                    field_values=[[config_app.session_id]],
                    order_by_field_names=['update_date'],
                    order_by_types=['DESC'],
                    limit=1)

    config_app.settings_id = int(records['id'].iloc[0])
    config_app.llm = records['llm'].iloc[0]
    config_app.temperature = float(records['temperature'].iloc[0])
    config_app.instructions = records['instructions'].iloc[0]

@print_func_name
def getFileType(file_name):

    if Path(file_name).suffix not in ['.docx', '.pdf']:
        return 'txt'
    return Path(file_name).suffix[1:]

@module
def getFileTypeIcon(input, output, session, file_type):
    
    @render.image()
    @print_func_name
    def icon():
        img: ImgData = {"src": str(Config.DIR_HOME / 'www' / 'assets' / f'{file_type}_icon.png'), 
                        "width": "100%"}
        return img
    
@print_func_name
def createVectorDBCollection(collection_name: str, replace_collection: bool=True):

    db = ChromaDB()
    if replace_collection: 
        db.create(collection_name=collection_name, delete_if_exists=True)
    else:
        db.get(collection_name=collection_name)

    return db

@print_func_name
def loadFilesToVectorDBCollection(collection_name: str, file_paths: list[tuple[Path, str]], replace_collection: bool=True, progress=None):

    def loadFiles(file_paths):

        docs = []

        progress_counter = 1

        for file_path, file_name in file_paths:
            
            if progress is not None:
                progress.set(progress_counter, f'Extracting file {progress_counter}')
                progress_counter += 1
            
            loader = getLoader(file_path=file_path)
            
            for doc in loader:
                doc.metadata = {**{'app_file_id': Path(file_path).stem, 'app_file_name': file_name}, **{k: str(v) for k, v in doc.metadata.items()}}
                docs.append(doc)

        if progress is not None:
            progress.set(progress_counter, 'Creating vector db')

        return docs

    docs = loadFiles(file_paths)

    db = createVectorDBCollection(collection_name=collection_name, replace_collection=replace_collection)
    db.add(docs=docs)
    
@print_func_name
def getVectorDBFiles(vector_db_collections_id):

    if not vector_db_collections_id: return [], {}

    vector_db_collection_records = selectFromDB(table_name='vector_db_collections', 
                                                field_names=['id', 'status'], 
                                                field_values=[[vector_db_collections_id], [vector_db_collections_status.ACTIVE.value]])
    
    if vector_db_collection_records.empty: return [], {}
    
    vector_db_collection_files_records = selectFromDB(table_name='vector_db_collection_files', 
                                                field_names=['vector_db_collections_id'], 
                                                field_values=[[vector_db_collections_id]])
    
    if vector_db_collection_files_records.empty: return [], {}

    uploaded_files_id_list = list(map(int, vector_db_collection_files_records['uploaded_files_id'].dropna().values))

    if uploaded_files_id_list: 

        uploaded_files_records = selectFromDB(table_name='uploaded_files',
                                            field_names=['id'],
                                            field_values=[uploaded_files_id_list])
        
        uploaded_files_records['type'] = vector_db_collections_type.UPLOADED_FILES.value
        
        uploaded_files_info = list(uploaded_files_records[['id', 'file_name', 'type']].values)

        uploaded_files_records['id'] = uploaded_files_records['id'].map(str)
        uploaded_files_records['title'] = uploaded_files_records['file_name']

        file_info = uploaded_files_records[['id', 'title']].set_index('id').T.to_dict()

        return uploaded_files_info, file_info 

    literature_id_list = list(vector_db_collection_files_records['literature_id'].dropna().values)

    if literature_id_list:
    
        return getLiteraturesFromDB(literature_id_list)
    
    return [], {}

@print_func_name
def getLiteraturesFromDB(literature_id_list):

    literature_records = selectFromDB(table_name='literature',
                                    field_names=['id'],
                                    field_values=[literature_id_list])
    
    literature_records['authors'] = literature_records['authors'].map(eval)
    literature_records['reference'] = literature_records.apply(lambda x: formatAPA(dict(x[['authors', 'title', 'year', 'journal', 'volume', 'issue', 'pages', 'doi', 'pmid']])), axis=1)
    literature_records['type'] = vector_db_collections_type.LITERATURE.value
    
    literature_info = list(literature_records[['id', 'reference', 'type']].values)

    literature_records['doi'] = literature_records['id']
    file_info = literature_records[['id', 'authors', 'title', 'journal', 'volume', 'issue', 'pages', 'year', 'doi']].set_index('id').T.to_dict()

    return literature_info, file_info

@print_func_name
def detachDocs(generated_files_id, vector_db_collections_id):

    current_time = datetime.now()
    
    updateDB(table_name='vector_db_collections', 
                update_fields=['status', 'update_date'],
                update_values=[vector_db_collections_status.DELETED.value, current_time],
                select_fields=['id'], 
                select_values=[[vector_db_collections_id]]) 
    
    records = selectFromDB(table_name='vector_db_collections',
                                field_names=['generated_files_id', 'status'], 
                                field_values=[[generated_files_id], [vector_db_collections_status.ACTIVE.value]])
    if records.empty:
        updateDB(table_name='generated_files', 
                update_fields=['ai_architecture', 'update_date'], 
                update_values=[generated_files_ai_architecture.BASE.value, current_time], 
                select_fields=['id'], 
                select_values=[[generated_files_id]])

    vector_db_collection_name = f'{Config.APP_NAME_AS_PREFIX}_collection_{vector_db_collections_id}'
    deleteCollection(vector_db_collection_name)

@print_func_name
def getGeneratedDocuments(email, session_id):

    valid_file_statuses = {e.value for e in generated_files_status} - {generated_files_status.DELETED.value}
    if email != '':
        records = selectFromDB(table_name='generated_files', 
                                field_names=['email', 'status'], 
                                field_values=[[email], valid_file_statuses],
                                order_by_field_names=['file_name'])
    else:
        records = selectFromDB(table_name='generated_files', 
                                field_names=['session', 'status'], 
                                field_values=[[session_id], valid_file_statuses],
                                order_by_field_names=['file_name'])
        
    
    if records.empty: return records

    records_settings = selectFromDB(table_name='settings',
                        field_names=['id'],
                        field_values=[list(map(int, records['settings_id'].unique()))])

    records = pd.merge(left=records, 
                        right=records_settings[['id', 'llm', 'temperature', 'instructions']], 
                        left_on='settings_id', right_on='id', how='left',
                        suffixes=[None, '_settings'])

    records_vector_db_collections = selectFromDB(table_name='vector_db_collections',
                        field_names=['generated_files_id', 'status'],
                        field_values=[list(map(int, records['id'].unique())), [vector_db_collections_status.ACTIVE.value]])

    if not records_vector_db_collections.empty:
        records = pd.merge(left=records, 
                            right=records_vector_db_collections[['id', 'generated_files_id', 'type']], 
                            left_on='id', right_on='generated_files_id', how='left',
                            suffixes=[None, '_vector_db_collections'])
    
    records_pivot = {}
    cols = list(records.columns[~records.columns.isin(['type', 'id_vector_db_collections', '_sa_instance_state'])])
    for i, row in records.iterrows():
        key = tuple(row[cols])
        records_pivot[key] = records_pivot.get(key, [None, None])
        if 'type' not in row.index: continue
        if not pd.isna(row['id_vector_db_collections']):
            if row['type'] == 'uploaded_files':
                records_pivot[key][0] = row['id_vector_db_collections']
            elif row['type'] == 'literature':
                records_pivot[key][1] = row['id_vector_db_collections']
    
    records_pivot = pd.DataFrame([list(k) + v for k, v in records_pivot.items()], columns = cols + ['vector_db_collections_id_uploaded_files', 'vector_db_collections_id_literature'])

    int_ = lambda x: 0 if pd.isna(x) else int(x)

    for col in ['vector_db_collections_id_uploaded_files', 'vector_db_collections_id_literature']:
        records_pivot[col] = records_pivot[col].apply(int_).astype('object')

    return records_pivot

@print_func_name
def getDocContent(file_id, attached_files=[], file_info={}):

    @print_func_name
    def processCitation(content, ref_list=[], used_files_info={}):

        content = formatCitations(content)

        content_tex = content
        
        ref_groups = re.findall(r'CITE\(([\w\W]+?)\)', content)
    
        refs_seen = set()
        d_ref = {}
        for refs in ref_groups:
            refs = re.sub(r'\),\ *CITE\(', ', ', refs)
            if refs in refs_seen: continue
            refs_seen.add(refs)
            ref_links = []
            for ref in refs.split(','):
                ref = ref.strip()
                if ref not in attached_references: 
                    logging.warning(f'{ref} not found in reference list, skipping...')
                    continue
                if ref in d_ref:
                    ref_links.append(d_ref[ref])
                    continue
                try:
                    d_ref[ref] = ref_list.index(attached_references[ref]) + 1
                except ValueError:
                    ref_list.append(attached_references[ref])
                    d_ref[ref] = len(ref_list)
                    used_files_info[ref] = file_info[ref]

                ref_links.append(d_ref[ref])
        
            ref_links = sorted(ref_links)
            if len(ref_links) > 2 and len(ref_links) == (ref_links[-1]-ref_links[0]+1):
                new_citation = f' [{ref_links[0]}-{ref_links[-1]}]'
            else:
                new_citation = f' [{', '.join(map(str, ref_links))}]'
            
            content = content.replace(f'[CITE({refs})]', new_citation)
            content_tex = content_tex.replace(f'[CITE({refs})]', f'\\cite{{{refs}}}')
    
        return content, content_tex, ref_list, used_files_info

    @print_func_name
    def extractContentFromOutline(d, content_md=[], content_docx=Document(), content_tex=[], ref_list=[], used_files_info={}, level=1):

        def latexLevels(level, header):

            match level:
                case 1:
                    return f'\\title{{{header}}}'
                case 2:
                    return f'\\section{{{header}}}'
                case 3:
                    return f'\\subsection{{{header}}}'
                case 4:
                    return f'\\subsubsection{{{header}}}'
                case 5:
                    return f'\\paragraph{{\\textbf{{{header}}}}}'
                case 6:
                    return f'\\paragraph{{\\textit{{{header}}}}}'

        if not isinstance(d, dict):
            for k, v in d:
                if k not in [ContentTypes.CONTENT_AI.value, ContentTypes.CONTENT_USER.value]: continue
                content_text, content_tex_text, ref_list, used_files_info = processCitation(v, ref_list, used_files_info)
                content_md.append(content_text)
                content_docx.add_paragraph(content_text)
                content_tex.append(content_tex_text)
        else:
            for k in d:
                if k != SpecialSectionTypes.CONTENT.value:
                    content_docx.add_heading(k, level=level)
                    content_md, content_docx, content_tex, ref_list, used_files_info = extractContentFromOutline(d[k], 
                                                                                                                 content_md + [f'{'#' * level} {k}'], 
                                                                                                                 content_docx, 
                                                                                                                 content_tex + [latexLevels(level, k)], 
                                                                                                                 ref_list, used_files_info, level+1)
                else:
                    content_md, content_docx, content_tex, ref_list, used_files_info = extractContentFromOutline(d[k], 
                                                                                                                 content_md, 
                                                                                                                 content_docx, 
                                                                                                                 content_tex, 
                                                                                                                 ref_list, used_files_info, level+1)

        return content_md, content_docx, content_tex, ref_list, used_files_info
    
    @print_func_name
    def convertToLatex(content):
        def formatLatex(text):

            """Escapes special characters in a string for LaTeX compatibility."""
            latex_special_chars = {
                '&': r'\&', '%': r'\%', '$': r'\$', '#': r'\#', '_': r'\_', '<': r'$<$', '>': r'$>$'
            }
            for char, replacement in latex_special_chars.items():
                text = text.replace(char, replacement)

            return text

        if not len(content): return ''
        return f"\\documentclass{{article}}\n\n{formatLatex(content[0])}\n\n\\begin{{document}}\n\n\\maketitle\n\n{formatLatex('\n'.join(content[1:]))}\n\n\\bibliographystyle{{plain}}\n\n\\bibliography{{bibliography}}\n\n\\end{{document}}"
    
    @print_func_name
    def getBibFormat(file_info):

        bib_text = ''
        for k, v in file_info.items():
            bib_ele = []
            for k1, v1 in v.items():
                if k1 == 'authors':
                    authors = [f'{author['first_name']} {author['last_name']}' for author in v1]
                    bib_ele.append(f'author = "{', '.join(authors)}"')
                elif k1 == 'pages':
                    bib_ele.append(f'{k1} = "{v1.replace('-', '--')}"')
                else:
                    bib_ele.append(f'{k1} = "{v1}"')
            bib_text += f'@article{{{k},\n\t{',\n\t'.join(bib_ele)}\n}}\n\n'

        return bib_text

    outline_file_path = Config.DIR_CONTENTS / f'outline_{file_id}.json'

    with open(outline_file_path) as fp:
        d_outline = json.load(fp)

    attached_references = {str(k): v for k, v, _ in attached_files}
    content_md, content_docx, content_tex, ref_list, used_files_info= extractContentFromOutline(d_outline)
    
    if attached_files:
        content_md.append('## References')
        content_docx.add_heading('References', level=2)
        for i, ref in enumerate(ref_list):
            content_md.append(f'\n[{i+1}] {ref}')
            content_docx.add_paragraph(f'[{i+1}] {ref}')
        bibs = getBibFormat(used_files_info)
    else:
        bibs = ''

    content_md = '\n'.join(content_md)
    content_tex = convertToLatex(content_tex)

    return content_md, content_docx, content_tex, bibs

@print_func_name
def uploadFiles(files, email='', session_id=''):

    for file in files:

        current_time = datetime.now()
        
        if email != '':
            records = selectFromDB(table_name='uploaded_files', 
                            field_names=['email', 'file_name', 'status'],
                            field_values=[[email], [file['name']], [uploaded_files_status.UPLOADED.value]])
        else:
            records = selectFromDB(table_name='uploaded_files', 
                            field_names=['session', 'file_name', 'status'],
                            field_values=[[session_id], [file['name']], [uploaded_files_status.UPLOADED.value]])
        
        if records.empty:

            ids = insertIntoDB(table_name='uploaded_files', 
                        field_names=['email', 'session', 'file_name', 'status', 'create_date', 'update_date'], 
                        field_values=[[email], [session_id], [file['name']], [uploaded_files_status.UPLOADED.value], [current_time], [current_time]])
            uploaded_file_id = int(ids[0])
            
        else:
            updateDB(table_name='uploaded_files', 
                    update_fields=['status', 'update_date'], 
                    update_values=[uploaded_files_status.UPLOADED.value, current_time], 
                    select_fields=['id'], 
                    select_values=[list(map(int, records.id.values))])
            uploaded_file_id = int(records.iloc[0].id)

        # ids = insertIntoDB(table_name='uploaded_files', 
        #                    field_names=['email', 'session', 'file_name', 'status', 'create_date', 'update_date'], 
        #                    field_values=[[email], [session_id], [file['name']], [uploaded_files_status.UPLOADED.value], [current_time], [current_time]])
        # uploaded_file_id = ids[0]
            
        dir_uploaded_files = Config.DIR_CONTENTS / 'uploaded_docs'
        dir_uploaded_files.mkdir(parents=False, exist_ok=True)

        with open(dir_uploaded_files / f'{uploaded_file_id}{Path(file['datapath']).suffix}', 'wb') as fp:
            with open(file['datapath'], 'rb') as fp_r:
                fp.write(fp_r.read())

@print_func_name
def unMarkdownText(text):

    from bs4 import BeautifulSoup
    from markdown import markdown

    html = markdown(text)
    return ''.join(BeautifulSoup(html).findAll(text=True))

@print_func_name
def formatCitations(text):
    '''
    Convert [CITE(abc), CITE(bcd), CITE(cde)] to [CITE(abc, bcd, cde)]
    '''
    pattern = r'\[(?:CITE\([^)]+\)(?:,\s*)?)+\]'
    
    def replace_func(match):
        # Extract all citations
        citations = re.findall(r'CITE\(([^)]+)\)', match.group(0))
        # Rebuild as single CITE with all arguments
        return f'[CITE({", ".join(citations)})]'
    
    return re.sub(pattern, replace_func, text)
    