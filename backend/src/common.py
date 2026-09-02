import json
import re
import logging
from docx import Document
from enum import Enum

from .utils import Config, print_func_name
from .vectordb import ChromaDB
from .db import selectFromDB, vector_db_collections_status, vector_db_collections_type
from .ai.tools.search_pubmed import formatAPA

# ---------------------------------------------------------------------------
class OutlineTAGs(Enum):
    INSTRUCTIONS_START = '[--instructions--]'
    INSTRUCTIONS_END = '[/--instructions--]'
    CONTENT = '[--content--]'

# ---------------------------------------------------------------------------
class SpecialSectionTypes(Enum):
    CONTENT = 'content'
    
# ---------------------------------------------------------------------------
class ContentTypes(Enum):
    IS_ABSTRACT = 'is_abstract'
    INSTRUCTIONS = 'instructions'
    CONTENT_USER = 'content_user'
    CONTENT_AI = 'content_ai'
    CONTENT_PRE_SUMMARY = 'content_pre_summary'
    CONCEPT_MAP = 'concept_map'
    KEYPHRASES = 'keyphrases'
    RETRIEVED_DOC_IDS = 'retrieved_doc_ids'

# ---------------------------------------------------------------------------
@print_func_name
def sanitizeContent(content):
    return re.sub(r' \~([^\~])', r' \\~\1', content)

# ---------------------------------------------------------------------------
@print_func_name
def formatCitations(text):
    '''
    Convert [CITE(abc), CITE(bcd), CITE(cde)] to [CITE(abc, bcd, cde)]
    '''
    pattern = r'\[(?:CITE\([^)]+\)(?:,\s*)?)+\]'
    
    def replace_func(match):
        # Extract all citations
        citations = re.findall(r'CITE\(([^)]+)\)', match.group(0))
        # Rebuild as single CITE with all arguments
        return f'[CITE({", ".join(citations)})]'
    
    return re.sub(pattern, replace_func, text)

# ---------------------------------------------------------------------------
@print_func_name
def rawCiteContent(content: str, 
                   attached_reference_list_from_content: list, 
                   attached_references_db: dict[str, str]) -> str:
    '''
    Replaces the formatted citation (e.g. [1] or [1, 2]) in the content back to the raw citations 
    (e.g. [CITE(<ref_id1>)] or [CITE(ref_id1), CITE(ref_id2)])
    '''
    
    matches = re.findall(r'<a\s+href="#:~:text=References"[^>]*>(.*?)<\/a>', content)

    ref_to_id = {v:k for k, v in attached_references_db.items()}
    ref_ids = {i+1:ref_to_id[ref] for i, ref in enumerate(attached_reference_list_from_content)}

    for m in matches:
        if '-' in m:
            start, end = m.split('-')
            raw_cite_text = ', '.join([f'CITE({ref_ids[n]})' for n in range(int(start), int(end)+1) if n in ref_ids])
        else:
            nums = m.split(', ')
            raw_cite_text = ', '.join([f'CITE({ref_ids[int(n)]})' for n in nums if n in ref_ids])
        
        content = re.sub(rf'<a\s+href="#:~:text=References"[^>]*>{m}<\/a>', raw_cite_text, content)

    return content

# ---------------------------------------------------------------------------
@print_func_name
def processCitation(content, attached_references_db, attached_reference_list_from_content=[], return_latex_style=False, enable_html_link_format=False):
    
    def formatCitationText(citation_text):
        if enable_html_link_format:
            return f'<a href="#:~:text=References">{citation_text}</a>'
        return citation_text
    
    def reEscape(text):
        return re.escape(text).replace('/', r'\/')

    content = formatCitations(content)

    if return_latex_style: content_tex = content
    
    ref_groups = re.findall(r'CITE\(([\w\W]+?)\)', content)

    refs_seen = set()
    d_ref = {}
    for refs in ref_groups:
        refs = re.sub(r'\),\ *CITE\(', ', ', refs)
        if refs in refs_seen: continue
        refs_seen.add(refs)
        ref_links = []
        for ref in refs.split(','):
            ref = ref.strip()
            if ref not in attached_references_db: 
                logging.warning(f'{ref} not found in reference list, skipping...')
                continue
            if ref in d_ref:
                ref_links.append(d_ref[ref])
                continue
            try:
                d_ref[ref] = attached_reference_list_from_content.index(attached_references_db[ref]) + 1
            except ValueError:
                attached_reference_list_from_content.append(attached_references_db[ref])
                d_ref[ref] = len(attached_reference_list_from_content)

            ref_links.append(d_ref[ref])
    
        ref_links = sorted(ref_links)
        if len(ref_links) > 2 and len(ref_links) == (ref_links[-1]-ref_links[0]+1):
            new_citation = formatCitationText(f'{ref_links[0]}-{ref_links[-1]}')
        else:
            new_citation = formatCitationText(', '.join(map(str, ref_links)))
        if new_citation == '':
            logging.warning(f'Invalid citation(s) or citation(s) not found in database: [CITE({refs})], removing ...')
            content = re.sub(rf' *\[CITE\({reEscape(refs)}\)\]', '', content)
            content_tex = re.sub(rf' *\[CITE\({reEscape(refs)}\)\]', '', content_tex)
        else:
            content = re.sub(rf' *\[CITE\({reEscape(refs)}\)\]', f' [{new_citation}]', content)
            if return_latex_style: content_tex = re.sub(rf' *\[CITE\({reEscape(refs)}\)\]', rf' \\cite{{{reEscape(refs)}}}', content_tex)
        
    if return_latex_style:
        return content, content_tex, attached_reference_list_from_content
    
    if 'CITE' in content: raise Warning('Some citations were not replaced, needs human review.')
    
    return content, attached_reference_list_from_content

# ---------------------------------------------------------------------------
@print_func_name
def getDocContent(file_id, vector_db_collections_id_uploaded_files, vector_db_collections_id_literature):

    @print_func_name
    def extractContentFromOutline(d, content_md=[], content_docx=Document(), content_tex=[], attached_reference_list_from_content=[], level=1):

        def latexLevels(level, header):

            match level:
                case 1:
                    return f'\\title{{{header}}}'
                case 2:
                    return f'\\section{{{header}}}'
                case 3:
                    return f'\\subsection{{{header}}}'
                case 4:
                    return f'\\subsubsection{{{header}}}'
                case 5:
                    return f'\\paragraph{{\\textbf{{{header}}}}}'
                case 6:
                    return f'\\paragraph{{\\textit{{{header}}}}}'

        if not isinstance(d, dict):
            for k, v in d:
                if k not in [ContentTypes.CONTENT_AI.value, ContentTypes.CONTENT_USER.value]: continue
                content_text, content_tex_text, attached_reference_list_from_content = processCitation(v, attached_references_db, attached_reference_list_from_content, return_latex_style=True)
                content_md.append(content_text)
                content_docx.add_paragraph(content_text)
                content_tex.append(content_tex_text)
        else:
            for k in d:
                if k != SpecialSectionTypes.CONTENT.value:
                    content_docx.add_heading(k, level=level)
                    content_md, content_docx, content_tex, attached_reference_list_from_content = extractContentFromOutline(d[k], 
                                                                                                content_md + [f'{'#' * level} {k}'], 
                                                                                                content_docx, 
                                                                                                content_tex + [latexLevels(level, k)], 
                                                                                                attached_reference_list_from_content, level+1)
                else:
                    content_md, content_docx, content_tex, attached_reference_list_from_content = extractContentFromOutline(d[k], 
                                                                                                content_md, 
                                                                                                content_docx, 
                                                                                                content_tex, 
                                                                                                attached_reference_list_from_content, level+1)

        return content_md, content_docx, content_tex, attached_reference_list_from_content
    
    @print_func_name
    def convertToLatex(content):
        def formatLatex(text):

            """Escapes special characters in a string for LaTeX compatibility."""
            latex_special_chars = {
                '&': r'\&', '%': r'\%', '$': r'\$', '#': r'\#', '_': r'\_', '<': r'$<$', '>': r'$>$'
            }
            for char, replacement in latex_special_chars.items():
                text = text.replace(char, replacement)

            return text

        if not len(content): return ''
        return f"\\documentclass{{article}}\n\n{formatLatex(content[0])}\n\n\\begin{{document}}\n\n\\maketitle\n\n{formatLatex('\n'.join(content[1:]))}\n\n\\bibliographystyle{{plain}}\n\n\\bibliography{{bibliography}}\n\n\\end{{document}}"
    
    @print_func_name
    def getBibFormat(file_info):

        bib_text = ''
        for k, v in file_info.items():
            bib_ele = []
            for k1, v1 in v.items():
                if k1 == 'authors':
                    authors = [f'{author['first_name']} {author['last_name']}' for author in v1]
                    bib_ele.append(f'author = "{', '.join(authors)}"')
                elif k1 == 'pages':
                    bib_ele.append(f'{k1} = "{v1.replace('-', '--')}"')
                else:
                    bib_ele.append(f'{k1} = "{v1}"')
            bib_text += f'@article{{{k},\n\t{',\n\t'.join(bib_ele)}\n}}\n\n'

        return bib_text

    outline_file_path = Config.DIR_CONTENTS / f'outline_{file_id}.json'

    with open(outline_file_path) as fp:
        d_outline = json.load(fp)

    attached_references_db, file_info = getAttachedRefs(vector_db_collections_id_uploaded_files, vector_db_collections_id_literature)

    attached_references_db = {str(k): v for k, v, _ in attached_references_db}
    content_md, content_docx, content_tex, attached_reference_list_from_content = extractContentFromOutline(d_outline)

    used_files_info = {ref: file_info[ref] for ref in attached_references_db}
    
    if attached_references_db:
        content_md.append('## References')
        content_docx.add_heading('References', level=2)
        for i, ref in enumerate(attached_reference_list_from_content):
            content_md.append(f'\n[{i+1}] {ref}')
            content_docx.add_paragraph(f'[{i+1}] {ref}')
        bibs = getBibFormat(used_files_info)
    else:
        bibs = ''

    content_md = '\n'.join(content_md)
    content_tex = convertToLatex(content_tex)

    return content_md, content_docx, content_tex, bibs

# ---------------------------------------------------------------------------
@print_func_name
def createVectorDBCollection(collection_name: str, replace_collection: bool=True):

    db = ChromaDB()
    if replace_collection: 
        db.create(collection_name=collection_name, delete_if_exists=True)
    else:
        db.get(collection_name=collection_name)

    return db

# ---------------------------------------------------------------------------
@print_func_name
def getLiteraturesFromDB(literature_id_list):

    literature_records = selectFromDB(table_name='literature',
                                    field_names=['id'],
                                    field_values=[literature_id_list])
    
    literature_records['authors'] = literature_records['authors'].map(eval)
    literature_records['reference'] = literature_records.apply(lambda x: formatAPA(dict(x[['authors', 'title', 'year', 'journal', 'volume', 'issue', 'pages', 'doi', 'pmid']])), axis=1)
    literature_records['type'] = vector_db_collections_type.LITERATURE.value
    
    literature_info = list(literature_records[['id', 'reference', 'type']].values)

    literature_records['doi'] = literature_records['id']
    file_info = literature_records[['id', 'authors', 'title', 'journal', 'volume', 'issue', 'pages', 'year', 'doi']].set_index('id').T.to_dict()

    return literature_info, file_info

# ---------------------------------------------------------------------------
@print_func_name
def getVectorDBFiles(vector_db_collections_id):

    if not vector_db_collections_id: return [], {}

    vector_db_collection_records = selectFromDB(table_name='vector_db_collections', 
                                                field_names=['id', 'status'], 
                                                field_values=[[vector_db_collections_id], [vector_db_collections_status.ACTIVE.value]])
    
    if vector_db_collection_records.empty: return [], {}
    
    vector_db_collection_files_records = selectFromDB(table_name='vector_db_collection_files', 
                                                field_names=['vector_db_collections_id'], 
                                                field_values=[[vector_db_collections_id]])
    
    if vector_db_collection_files_records.empty: return [], {}

    uploaded_files_id_list = list(map(int, vector_db_collection_files_records['uploaded_files_id'].dropna().values))

    if uploaded_files_id_list: 

        uploaded_files_records = selectFromDB(table_name='uploaded_files',
                                            field_names=['id'],
                                            field_values=[uploaded_files_id_list])
        
        uploaded_files_records['type'] = vector_db_collections_type.UPLOADED_FILES.value
        
        uploaded_files_info = list(uploaded_files_records[['id', 'file_name', 'type']].values)

        uploaded_files_records['id'] = uploaded_files_records['id'].map(str)
        uploaded_files_records['title'] = uploaded_files_records['file_name']

        file_info = uploaded_files_records[['id', 'title']].set_index('id').T.to_dict()

        return uploaded_files_info, file_info 

    literature_id_list = list(vector_db_collection_files_records['literature_id'].dropna().values)

    if literature_id_list:
    
        return getLiteraturesFromDB(literature_id_list)
    
    return [], {}

# ---------------------------------------------------------------------------
@print_func_name
def getAttachedRefs(vector_db_collections_id_uploaded_files, vector_db_collections_id_literature):

    files_attached, file_info_attached = getVectorDBFiles(vector_db_collections_id_uploaded_files)
    files_lit, file_info_lit = getVectorDBFiles(vector_db_collections_id_literature)
    
    return files_attached + files_lit, file_info_attached | file_info_lit