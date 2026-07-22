from __future__ import annotations

import asyncio
from copy import deepcopy
from datetime import datetime
from io import BytesIO
import json
import logging
from pathlib import Path
import re
import shutil
import threading
import zipfile
from typing import Any
from uuid import uuid4
from typing import Literal
from pydantic import BaseModel, Field

from fastapi import HTTPException, Request, UploadFile
from fastapi.responses import RedirectResponse, Response, StreamingResponse
from fastapi.encoders import jsonable_encoder

from src.auth import (
    authenticate_azure_session,
    authenticate_user,
    azure_callback as auth_azure_callback,
    azure_login as auth_azure_login,
    azure_status as auth_azure_status,
    change_password as auth_change_password,
    create_account as auth_create_account,
    create_guest_auth_payload as _create_guest_auth_payload,
    credential_by_id as _credential_by_id,
    reset_password as auth_reset_password,
    send_password_reset_code,
    validate_email as _validate_email,
    validate_maintainer_access,
    verify_reset_code as auth_verify_reset_code,
)
from src.utils import Config


logger = logging.getLogger(f"{Config.APP_NAME_AS_PREFIX}.api")

class AIRequestBase(BaseModel):
    model_name: str | None = None
    temperature: float = 0
    instructions: str = ""


class ContentRequest(AIRequestBase):
    architecture_type: Literal["base", "rag", "graphrag"] = "base"
    collection_name: str = ""
    collection_name_lit_search: str = ""
    content_pre_summary: str = ""
    current_section: str
    content_specific_instructions: str = ""
    keyphrases: list[str] = Field(default_factory=list)
    rag_context: str = ""
    graphrag_context: dict[str, Any] = Field(default_factory=dict)
    literature_list: list[dict[str, Any]] = Field(default_factory=list)
    references: list[dict[str, Any]] = Field(default_factory=list)
    is_abstract: bool = False
    concept_map: dict[str, Any] = Field(default_factory=dict)


class OutlineFormatRequest(AIRequestBase):
    outline_unstructured: str
    query: str = ""


class GeneratedFileRequest(BaseModel):
    email: str | None = None
    session: str
    settings_id: int | None = None
    file_name: str
    ai_architecture: Literal["base", "rag", "graphrag"] = "base"
    replace: bool = False


class GeneratedFileUpdateRequest(BaseModel):
    file_name: str
    replace: bool = False


class UploadedFileUpdateRequest(BaseModel):
    email: str | None = None
    session: str | None = None
    file_name: str


class GeneratedFileLiteratureSearchRequest(BaseModel):
    email: str | None = None
    session: str | None = None


class GeneratedFileAttachUploadedFilesRequest(BaseModel):
    uploaded_file_ids: list[int] = Field(default_factory=list)
    email: str | None = None
    session: str | None = None
    mode: Literal["ask", "append", "replace"] = "ask"


class GeneratedFileGenerateRequest(AIRequestBase):
    outline: str
    email: str | None = None
    session: str | None = None
    mode: Literal["remaining", "restart"] = "remaining"
    architecture_type: Literal["base", "rag", "graphrag"] | None = None
    collection_name: str = ""
    collection_name_lit_search: str = ""
    attached_references_db: dict[str, str] = Field(default_factory=dict)


class GeneratedFileParagraphUpdateRequest(AIRequestBase):
    section_path: list[str] = Field(default_factory=list)
    section_heading: str = ""
    paragraph_index: int = Field(ge=0)
    raw_paragraph: str
    action: Literal["Expand", "Rephrase", "Remove"]
    email: str | None = None
    session: str | None = None


class GeneratedFileSectionContentUpdateRequest(BaseModel):
    section_path: list[str] = Field(default_factory=list)
    section_heading: str = ""
    section_content: str
    attached_reference_list_from_content: list[str] = Field(default_factory=list)
    email: str | None = None
    session: str | None = None


DownloadFormat = Literal["md", "docx", "latex"]
LogSortField = Literal["date", "status", "message"]
SortDirection = Literal["asc", "desc"]


class SettingsUpdateRequest(BaseModel):
    llm: str
    temperature: float = Field(ge=0, le=2)
    instructions: str = ""


class MaintainerLogClearRequest(BaseModel):
    email: str
    session: str
    entry_ids: list[int] = Field(default_factory=list)


class UploadedOutlineTemplateUpdateRequest(BaseModel):
    credentials_id: int
    outline: str


class UploadedOutlineTemplateRenameRequest(BaseModel):
    credentials_id: int
    file_name: str


def _required_default_model() -> str:
    model_name = Config.env_config.get("DEFAULT_AI_MODEL")
    if not model_name:
        raise HTTPException(
            status_code=503,
            detail="DEFAULT_AI_MODEL is not configured. Add DEFAULT_AI_MODEL to backend/.env and restart the backend.",
        )
    return model_name


def _model_name(model_name: str | None) -> str:
    return model_name or _required_default_model()


def _exception_chain(exp: Exception) -> list[BaseException]:
    chain: list[BaseException] = []
    current: BaseException | None = exp
    while current is not None and current not in chain:
        chain.append(current)
        current = current.__cause__ or current.__context__
    return chain


def _is_database_error(exp: Exception) -> bool:
    database_tokens = ("sqlalchemy", "psycopg", "database", "operationalerror", "dbapierror")
    for current in _exception_chain(exp):
        current_type = type(current)
        type_text = f"{current_type.__module__}.{current_type.__name__}".lower()
        if any(token in type_text for token in database_tokens):
            return True
    return False


def _friendly_error_detail(action: str, exp: Exception) -> tuple[int, str]:
    if _is_database_error(exp):
        return 503, "The database is temporarily unavailable. Please try again in a moment."
    return 500, f"Unable to {action}. Please try again."


def _api_error(action: str, exp: Exception) -> HTTPException:
    logger.exception("Unable to %s", action)
    status_code, detail = _friendly_error_detail(action, exp)
    return HTTPException(status_code=status_code, detail=detail)


_LOG_LINE_PATTERN = re.compile(r"^(?P<date>\d{4}-\d{2}-\d{2} \d{2}:\d{2}:\d{2}) \[(?P<status>[A-Z]+)\] (?P<message>.*)$")


def _parse_log_date(value: str | None, end_of_day: bool = False) -> datetime | None:
    if not value:
        return None
    text = value.strip()
    for date_format in ("%Y-%m-%dT%H:%M:%S", "%Y-%m-%d %H:%M:%S", "%Y-%m-%d"):
        try:
            parsed = datetime.strptime(text, date_format)
            if date_format == "%Y-%m-%d" and end_of_day:
                return parsed.replace(hour=23, minute=59, second=59)
            return parsed
        except ValueError:
            continue
    raise HTTPException(status_code=400, detail="Use YYYY-MM-DD for log date filters.")


def _log_entries() -> list[dict[str, Any]]:
    log_file_path = Config.DIR_HOME / "logs" / "app.log"
    if not log_file_path.exists():
        return []

    entries: list[dict[str, Any]] = []
    entry_id = 0
    with log_file_path.open("r", encoding="utf-8", errors="replace") as fp:
        for line in fp:
            clean_line = line.rstrip("\n")
            match = _LOG_LINE_PATTERN.match(clean_line)
            if match:
                entries.append(
                    {
                        "id": entry_id,
                        "date": match.group("date"),
                        "status": match.group("status"),
                        "message": match.group("message"),
                        "_lines": [line],
                    }
                )
                entry_id += 1
                continue

            if entries and clean_line.strip():
                entries[-1]["message"] = f"{entries[-1]['message']}\n{clean_line}"
                entries[-1]["_lines"].append(line)

    return entries


def _public_log_entry(entry: dict[str, Any]) -> dict[str, Any]:
    return {key: value for key, value in entry.items() if not key.startswith("_")}


def _clear_log_entries(entry_ids: list[int]) -> int:
    ids_to_remove = {int(entry_id) for entry_id in entry_ids}
    if not ids_to_remove:
        raise HTTPException(status_code=400, detail="Select at least one log entry to clear.")

    log_file_path = Config.DIR_HOME / "logs" / "app.log"
    entries = _log_entries()
    removed_count = sum(1 for entry in entries if int(entry["id"]) in ids_to_remove)
    if removed_count == 0:
        raise HTTPException(status_code=404, detail="Selected log entries were not found.")

    log_file_path.parent.mkdir(parents=True, exist_ok=True)
    with log_file_path.open("w", encoding="utf-8") as fp:
        for entry in entries:
            if int(entry["id"]) in ids_to_remove:
                continue
            for line in entry.get("_lines", []):
                fp.write(line if line.endswith("\n") else f"{line}\n")

    return removed_count


def _service_health(status: str, label: str, message: str) -> dict[str, str]:
    return {
        "status": status,
        "label": label,
        "message": message,
    }


def _check_ai_model_health() -> dict[str, str]:
    model_name = Config.env_config.get("DEFAULT_AI_MODEL")
    missing = [
        name
        for name, value in (
            ("DEFAULT_AI_MODEL", model_name),
            ("AI_BASE_URL", Config.env_config.get("AI_BASE_URL")),
            ("AI_API_KEY", Config.env_config.get("AI_API_KEY")),
        )
        if not value
    ]
    if missing:
        return _service_health("error", "AI Model", f"Missing configuration: {', '.join(missing)}.")

    try:
        from src.ai.llms import getAIModel

        getAIModel(model_name=str(model_name), temperature=0)
        return _service_health("ok", "AI Model", f"{model_name} is configured.")
    except Exception:
        logger.exception("AI model health check failed")
        return _service_health("error", "AI Model", "The AI model client could not be initialized.")


def _check_chroma_health() -> dict[str, str]:
    host = Config.env_config.get("CHROMA_HOST")
    port = Config.env_config.get("CHROMA_PORT")
    if not host or not port:
        return _service_health("error", "Chroma DB", "Chroma host or port is not configured.")

    try:
        from chromadb import HttpClient

        client = HttpClient(host=host, port=port)
        client.heartbeat()
        return _service_health("ok", "Chroma DB", f"Chroma is reachable at {host}:{port}.")
    except Exception:
        logger.exception("Chroma DB health check failed")
        return _service_health("error", "Chroma DB", "Chroma DB is not reachable.")


def _check_postgres_health() -> dict[str, str]:
    try:
        import sqlalchemy as sa
        from src import db

        if not hasattr(db, "engine"):
            return _service_health("error", "Postgres", "Postgres engine is not initialized.")

        with db.engine.connect() as connection:
            connection.execute(sa.text("SELECT 1"))
        return _service_health("ok", "Postgres", "Postgres database is reachable.")
    except Exception:
        logger.exception("Postgres health check failed")
        return _service_health("error", "Postgres", "Postgres database is not reachable.")


def _jsonable(value: Any) -> Any:
    try:
        return jsonable_encoder(value)
    except Exception:
        if isinstance(value, dict):
            return {str(k): _jsonable(v) for k, v in value.items()}
        if isinstance(value, (list, tuple, set)):
            return [_jsonable(v) for v in value]
        if isinstance(value, datetime):
            return value.isoformat()
        return repr(value)


def _content_state(request: ContentRequest) -> dict[str, Any]:
    return {
        "content_pre_summary": request.content_pre_summary,
        "current_section": request.current_section,
        "content_specific_instructions": request.content_specific_instructions,
        "keyphrases": request.keyphrases,
        "rag_context": request.rag_context,
        "graphrag_context": request.graphrag_context,
        "literature_list": request.literature_list,
        "steps": [],
        "content": "",
        "content_summary": "",
        "references": request.references,
        "is_abstract": request.is_abstract,
        "concept_map": request.concept_map,
    }


def _outline_state(query: str = "", outline_unstructured: str = "") -> dict[str, Any]:
    return {
        "query": query,
        "steps": [],
        "reference_summary": "",
        "content": "",
        "outline_unstructured": outline_unstructured,
    }


def _records_from_dataframe(df: Any) -> list[dict[str, Any]]:
    df = df.drop(columns=["_sa_instance_state"], errors="ignore")
    return _jsonable(df.to_dict(orient="records"))


def _settings_record(record: dict[str, Any]) -> dict[str, Any]:
    return {
        "id": record.get("id"),
        "email": record.get("email"),
        "session": record.get("session"),
        "llm": record.get("llm"),
        "temperature": record.get("temperature"),
        "instructions": record.get("instructions") or "",
        "create_date": record.get("create_date"),
        "update_date": record.get("update_date"),
    }

def _llm_options() -> list[dict[str, str]]:
    from src.ai.llms import extractAvailableLLMs

    options = []
    default_model = _required_default_model()
    available_llms = extractAvailableLLMs()

    if isinstance(available_llms, dict):
        for provider, models in available_llms.items():
            if isinstance(models, dict):
                for value, label in models.items():
                    model_value = label if value == "Uncategorized" else value
                    options.append({"provider": provider, "value": str(model_value), "label": str(label)})
            else:
                options.append({"provider": provider, "value": str(models), "label": str(models)})
    else:
        for model in available_llms:
            options.append({"provider": "Available", "value": str(model), "label": str(model)})

    if not any(option["value"] == default_model for option in options):
        options.insert(0, {"provider": "Default", "value": default_model, "label": default_model})

    return options


def _generated_file_record(record: dict[str, Any]) -> dict[str, Any]:
    return {
        "id": record.get("id"),
        "email": record.get("email"),
        "session": record.get("session"),
        "file_name": record.get("file_name"),
        "status": record.get("status"),
        "settings_id": record.get("settings_id"),
        "ai_architecture": record.get("ai_architecture"),
        "create_date": record.get("create_date"),
        "update_date": record.get("update_date"),
    }


def _generated_document(record: dict[str, Any]) -> dict[str, Any]:
    return {
        "id": record.get("id"),
        "name": record.get("file_name"),
        "date": record.get("update_date") or record.get("create_date"),
        "status": record.get("status"),
        "session": record.get("session"),
        "settings_id": record.get("settings_id"),
        "ai_architecture": record.get("ai_architecture"),
    }


def _settings_for_generated_file(record: dict[str, Any]) -> dict[str, Any] | None:
    settings_id = record.get("settings_id")
    if settings_id is None:
        return None

    from src import db

    df = db.selectFromDB(
        table_name="settings",
        field_names=["id"],
        field_values=[[int(settings_id)]],
        limit=1,
    )
    records = _records_from_dataframe(df)
    return _settings_record(records[0]) if records else None


def _settings_summary(settings: dict[str, Any] | None) -> str:
    if not settings:
        return "Session defaults"

    llm = str(settings.get("llm") or "Default model")
    temperature = settings.get("temperature")
    if temperature is None:
        return llm
    return f"{llm} | Temp {temperature}"


def _generated_document_detail(record: dict[str, Any]) -> dict[str, Any]:
    settings = _settings_for_generated_file(record)
    attached_documents = _attached_uploaded_documents(int(record["id"])) if record.get("id") is not None else []
    return {
        **_generated_document(record),
        "file_name": record.get("file_name"),
        "last_modified": record.get("update_date") or record.get("create_date"),
        "attached_documents": attached_documents,
        "attached_documents_count": len(attached_documents),
        "settings": settings,
        "settings_summary": _settings_summary(settings),
    }


def _generated_records_for_owner(
    email: str | None = None,
    session: str | None = None,
    limit: int = 50,
) -> list[dict[str, Any]]:
    from src import db

    _, _, owner_field, owner_value = _owner_identity(email=email, session=session)
    active_statuses = [
        status.value for status in db.generated_files_status if status != db.generated_files_status.DELETED
    ]
    df = db.selectFromDB(
        table_name="generated_files",
        field_names=[owner_field, "status"],
        field_values=[[owner_value], active_statuses],
        order_by_field_names=["update_date"],
        order_by_types=["DESC"],
        limit=limit,
    )
    return _records_from_dataframe(df)


def _generated_file_by_id(
    generated_file_id: int,
    email: str | None = None,
    session: str | None = None,
) -> dict[str, Any]:
    from src import db

    _, _, owner_field, owner_value = _owner_identity(email=email, session=session)
    field_names = ["id", owner_field]
    field_values = [[generated_file_id], [owner_value]]

    df = db.selectFromDB(
        table_name="generated_files",
        field_names=field_names,
        field_values=field_values,
        limit=1,
    )
    records = _records_from_dataframe(df)
    if not records:
        raise HTTPException(status_code=404, detail="Generated file was not found.")

    return records[0]


def _duplicate_generated_file_for_owner(
    file_name: str,
    email: str | None = None,
    session: str | None = None,
    exclude_id: int | None = None,
) -> dict[str, Any] | None:
    normalized_name = file_name.strip().casefold()
    if not normalized_name:
        return None

    records = _generated_records_for_owner(email=email, session=session, limit=1000)
    for record in records:
        if exclude_id is not None and str(record.get("id")) == str(exclude_id):
            continue
        if str(record.get("file_name") or "").strip().casefold() == normalized_name:
            return record
    return None


def _mark_generated_file_replaced(record: dict[str, Any], now: datetime | None = None) -> None:
    from src import db

    generated_file_id = int(record["id"])
    now = now or datetime.now()
    db.updateDB(
        table_name="generated_files",
        update_fields=["status", "update_date"],
        update_values=[db.generated_files_status.DELETED.value, now],
        select_fields=["id"],
        select_values=[[generated_file_id]],
    )

    active_collections = db.selectFromDB(
        table_name="vector_db_collections",
        field_names=["generated_files_id", "status"],
        field_values=[[generated_file_id], [db.vector_db_collections_status.ACTIVE.value]],
    )
    for collection in _records_from_dataframe(active_collections):
        _delete_vector_collection_record(collection, now)


def _uploaded_file_type(file_name: str | None) -> str:
    suffix = Path(file_name or "").suffix.lower().lstrip(".")
    if suffix in {"doc", "docx"}:
        return "doc"
    if suffix == "pdf":
        return "pdf"
    return suffix or "file"


OUTLINE_REFERENCE_ALLOWED_SUFFIXES = {".txt", ".md", ".markdown", ".csv", ".tsv", ".docx", ".pdf"}
OUTLINE_REFERENCE_MAX_BYTES = 10 * 1024 * 1024
OUTLINE_IMPORT_ALLOWED_SUFFIXES = {".md", ".docx"}
OUTLINE_IMPORT_MAX_BYTES = 10 * 1024 * 1024
MARKDOWN_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
NUMBERED_OUTLINE_RE = re.compile(r"^(\d+(?:\.\d+)*)(?:[.)])?\s+(.+?)\s*$")
ROMAN_OUTLINE_RE = re.compile(r"^([IVXLCDM]+)[.)]\s+(.+?)\s*$", re.IGNORECASE)
ALPHA_OUTLINE_RE = re.compile(r"^([A-Z])[.)]\s+(.+?)\s*$")
BULLET_OUTLINE_RE = re.compile(r"^[-*+]\s+(.+?)\s*$")


def _decode_text_document(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore")


def _extract_docx_text(content: bytes) -> str:
    from docx import Document

    document = Document(BytesIO(content))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n\n".join(parts).strip()


def _clean_outline_heading(text: str) -> str:
    heading = re.sub(r"\s+#+\s*$", "", text).strip()
    heading = re.sub(r"^[*_`]+|[*_`]+$", "", heading).strip()
    return heading


def _line_heading_level(line: str) -> int | None:
    match = MARKDOWN_HEADING_RE.match(line.strip())
    return len(match.group(1)) if match else None


def _ensure_outline_content_tags(lines: list[str]) -> list[str]:
    from src.common import OutlineTAGs

    content_tag = OutlineTAGs.CONTENT.value
    heading_positions = [(index, _line_heading_level(line)) for index, line in enumerate(lines)]
    heading_positions = [(index, level) for index, level in heading_positions if level is not None]
    if not heading_positions:
        return lines

    insert_positions = set()
    for position_index, (line_index, level) in enumerate(heading_positions):
        next_heading = heading_positions[position_index + 1] if position_index + 1 < len(heading_positions) else None
        next_heading_index = next_heading[0] if next_heading else len(lines)
        next_heading_level = next_heading[1] if next_heading else None
        is_leaf = next_heading_level is None or next_heading_level <= level
        if not is_leaf:
            continue

        section_lines = lines[line_index + 1 : next_heading_index]
        if any(line.strip() == content_tag for line in section_lines):
            continue
        insert_positions.add(next_heading_index)

    normalized_lines = []
    for index, line in enumerate(lines):
        if index in insert_positions:
            normalized_lines.extend(["", content_tag, ""])
        normalized_lines.append(line)
    if len(lines) in insert_positions:
        normalized_lines.extend(["", content_tag])

    return normalized_lines


def _normalize_outline_lines(lines: list[str]) -> str:
    from src.manage_outline import getRawOutline, processOutline

    outline_text = "\n".join(_ensure_outline_content_tags(lines)).strip()
    if not outline_text:
        raise HTTPException(status_code=400, detail="No outline headings were found in that file.")

    try:
        processed_outline = processOutline(outline_text)
        return "\n".join(getRawOutline(processed_outline, raw_outline=[])).strip()
    except Exception:
        logger.exception("Unable to convert uploaded outline into app outline format.")
        raise HTTPException(
            status_code=400,
            detail="We found outline text, but could not convert it into the app outline format. Check the heading levels and try again.",
        )


def _markdown_outline_lines(text: str) -> list[str]:
    from src.common import OutlineTAGs

    lines = []
    inside_instructions = False
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        if line == OutlineTAGs.INSTRUCTIONS_START.value:
            inside_instructions = True
            lines.append(line)
            continue
        if line == OutlineTAGs.INSTRUCTIONS_END.value:
            inside_instructions = False
            lines.append(line)
            continue
        if inside_instructions or line == OutlineTAGs.CONTENT.value:
            lines.append(line)
            continue

        heading_match = MARKDOWN_HEADING_RE.match(line)
        if heading_match:
            heading = _clean_outline_heading(heading_match.group(2))
            if heading:
                lines.append(f"{heading_match.group(1)} {heading}")

    return lines


def _inferred_outline_lines(text: str, file_name: str | None = None) -> list[str]:
    meaningful_lines = [line.strip() for line in text.splitlines() if line.strip()]
    inferred_lines = []
    has_top_level = False
    fallback_title = _clean_outline_heading(Path(file_name or "Uploaded outline").stem.replace("_", " ").replace("-", " "))

    for line in meaningful_lines:
        line = re.sub(r"\s+", " ", line).strip()
        if not line or len(line) > 220:
            continue

        numbered_match = NUMBERED_OUTLINE_RE.match(line)
        roman_match = ROMAN_OUTLINE_RE.match(line)
        alpha_match = ALPHA_OUTLINE_RE.match(line)
        bullet_match = BULLET_OUTLINE_RE.match(line)

        if numbered_match:
            level = min(6, max(2, numbered_match.group(1).count(".") + 2))
            heading = numbered_match.group(2)
        elif roman_match:
            level = 2
            heading = roman_match.group(2)
        elif alpha_match:
            level = 3
            heading = alpha_match.group(2)
        elif bullet_match:
            level = 2
            heading = bullet_match.group(1)
        elif not has_top_level:
            level = 1
            heading = line
        else:
            continue

        heading = _clean_outline_heading(heading)
        if not heading:
            continue
        if level == 1:
            has_top_level = True
        inferred_lines.append(f"{'#' * level} {heading}")

    if inferred_lines and _line_heading_level(inferred_lines[0]) != 1:
        inferred_lines.insert(0, f"# Title: {fallback_title}")

    return inferred_lines


def _docx_outline_lines(content: bytes, file_name: str | None = None) -> list[str]:
    from docx import Document

    document = Document(BytesIO(content))
    lines = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (paragraph.style.name or "").strip().lower() if paragraph.style else ""
        heading_match = re.search(r"heading\s*([1-6])", style_name)
        if heading_match:
            level = int(heading_match.group(1))
            lines.append(f"{'#' * level} {_clean_outline_heading(text)}")

    if lines:
        if _line_heading_level(lines[0]) != 1:
            fallback_title = _clean_outline_heading(Path(file_name or "Uploaded outline").stem.replace("_", " ").replace("-", " "))
            lines.insert(0, f"# Title: {fallback_title}")
        return lines

    return _inferred_outline_lines(_extract_docx_text(content), file_name=file_name)


def _outline_from_uploaded_file(file_name: str | None, content: bytes) -> str:
    suffix = Path(file_name or "").suffix.lower()
    if suffix not in OUTLINE_IMPORT_ALLOWED_SUFFIXES:
        raise HTTPException(status_code=400, detail="Upload a Markdown (.md) or Word (.docx) outline file.")
    if len(content) > OUTLINE_IMPORT_MAX_BYTES:
        raise HTTPException(status_code=400, detail="Outline file must be 10 MB or smaller.")

    try:
        if suffix in {".md", ".markdown"}:
            text = _decode_text_document(content)
            lines = _markdown_outline_lines(text) or _inferred_outline_lines(text, file_name=file_name)
        else:
            lines = _docx_outline_lines(content, file_name=file_name)
    except HTTPException:
        raise
    except Exception:
        logger.exception("Unable to read uploaded outline file %s", file_name)
        raise HTTPException(status_code=400, detail="We could not read that outline file. Try a Markdown or DOCX file.")

    return _normalize_outline_lines(lines)


def _extract_pdf_text(content: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(content))
    return "\n\n".join(page.extract_text() or "" for page in reader.pages).strip()


def _query_file_path(generated_file_id: int) -> Path:
    return Config.DIR_CONTENTS / f"query_{generated_file_id}.md"


def _query_reference_dir(generated_file_id: int) -> Path:
    return Config.DIR_CONTENTS / f"ref_files_for_query_{generated_file_id}"


def _query_details_for_generated_file(generated_file_id: int) -> dict[str, Any]:
    query_path = _query_file_path(generated_file_id)
    reference_dir = _query_reference_dir(generated_file_id)
    reference_files = []
    if reference_dir.exists():
        reference_files = [
            {
                "name": path.name,
                "size": path.stat().st_size,
                "last_modified": datetime.fromtimestamp(path.stat().st_mtime).isoformat(),
                "saved_reference": True,
            }
            for path in sorted(reference_dir.iterdir())
            if path.is_file()
        ]

    return {
        "content": query_path.read_text(encoding="utf-8") if query_path.exists() else "",
        "file_name": query_path.name,
        "reference_files": reference_files,
    }


def _save_query_text(generated_file_id: int | None, query: str) -> Path | None:
    if generated_file_id is None:
        return None

    query_path = _query_file_path(generated_file_id)
    query_path.parent.mkdir(parents=True, exist_ok=True)
    query_path.write_text(query.strip(), encoding="utf-8")
    return query_path


def _save_query_reference_file(generated_file_id: int | None, upload: UploadFile, content: bytes) -> Path | None:
    if generated_file_id is None:
        return None

    reference_dir = _query_reference_dir(generated_file_id)
    reference_dir.mkdir(parents=True, exist_ok=True)
    file_name = _safe_uploaded_file_name(upload.filename)
    destination = reference_dir / file_name
    destination.write_bytes(content)
    return destination


async def _outline_reference_document_details(upload: UploadFile | None, generated_file_id: int | None = None) -> str:
    if upload is None or not upload.filename:
        return ""

    suffix = Path(upload.filename).suffix.lower()
    if suffix not in OUTLINE_REFERENCE_ALLOWED_SUFFIXES:
        raise HTTPException(
            status_code=400,
            detail="Use a supported reference document format: .txt, .md, .csv, .tsv, .docx, or .pdf.",
        )

    content = await upload.read()
    if len(content) > OUTLINE_REFERENCE_MAX_BYTES:
        raise HTTPException(status_code=400, detail="Reference document must be 10 MB or smaller.")

    _save_query_reference_file(generated_file_id, upload, content)

    try:
        if suffix in {".txt", ".md", ".markdown", ".csv", ".tsv"}:
            details = _decode_text_document(content).strip()
        elif suffix == ".docx":
            details = _extract_docx_text(content)
        elif suffix == ".pdf":
            details = _extract_pdf_text(content)
        else:
            details = ""
    except Exception:
        logger.exception("Unable to extract outline reference document text from %s", upload.filename)
        raise HTTPException(
            status_code=400,
            detail="We could not read text from that reference document. Try a text, DOCX, or searchable PDF file.",
        )

    if not details.strip():
        raise HTTPException(status_code=400, detail="No readable text was found in that reference document.")

    return details


async def _outline_reference_documents_details(uploads: list[UploadFile], generated_file_id: int | None = None) -> str:
    detail_blocks = []
    for upload in uploads:
        details = await _outline_reference_document_details(upload, generated_file_id)
        if details.strip():
            detail_blocks.append(f"# Reference file: {_safe_uploaded_file_name(upload.filename)}\n\n{details.strip()}")
    return "\n\n".join(detail_blocks).strip()


async def _outline_payload_from_request(request: Request) -> dict[str, Any]:
    def parse_generated_file_id(value: Any) -> int | None:
        if value in (None, ""):
            return None
        try:
            generated_file_id = int(value)
        except (TypeError, ValueError):
            raise HTTPException(status_code=400, detail="Generated file id must be a valid integer.")
        if generated_file_id <= 0:
            raise HTTPException(status_code=400, detail="Generated file id must be a valid integer.")
        return generated_file_id

    content_type = request.headers.get("content-type", "")
    if "multipart/form-data" in content_type or "application/x-www-form-urlencoded" in content_type:
        form = await request.form()
        reference_documents = [
            upload
            for upload in [*form.getlist("reference_documents"), *form.getlist("reference_document")]
            if hasattr(upload, "filename") and hasattr(upload, "read")
        ]

        try:
            temperature = float(form.get("temperature") or 0)
        except (TypeError, ValueError):
            raise HTTPException(status_code=400, detail="Temperature must be a valid number.")

        return {
            "query": str(form.get("query") or "").strip(),
            "model_name": str(form.get("model_name") or "").strip() or None,
            "temperature": temperature,
            "instructions": str(form.get("instructions") or ""),
            "reference_documents": reference_documents,
            "generated_file_id": parse_generated_file_id(form.get("generated_file_id")),
        }

    try:
        payload = await request.json()
    except Exception:
        payload = {}

    if not isinstance(payload, dict):
        raise HTTPException(status_code=400, detail="Outline request was not formatted correctly.")

    try:
        temperature = float(payload.get("temperature") or 0)
    except (TypeError, ValueError):
        raise HTTPException(status_code=400, detail="Temperature must be a valid number.")

    return {
        "query": str(payload.get("query") or "").strip(),
        "model_name": str(payload.get("model_name") or "").strip() or None,
        "temperature": temperature,
        "instructions": str(payload.get("instructions") or ""),
        "reference_documents": [],
        "generated_file_id": parse_generated_file_id(payload.get("generated_file_id")),
    }

def _uploaded_file_record(record: dict[str, Any]) -> dict[str, Any]:
    return {
        "id": record.get("id"),
        "email": record.get("email"),
        "session": record.get("session"),
        "file_name": record.get("file_name"),
        "status": record.get("status"),
        "create_date": record.get("create_date"),
        "update_date": record.get("update_date"),
    }


def _uploaded_document(record: dict[str, Any]) -> dict[str, Any]:
    return {
        "id": record.get("id"),
        "name": record.get("file_name"),
        "date": record.get("update_date") or record.get("create_date"),
        "type": _uploaded_file_type(record.get("file_name")),
        "status": record.get("status"),
        "session": record.get("session"),
    }


def _owner_identity(email: str | None = None, session: str | None = None) -> tuple[str, str, str, str]:
    normalized_email = _validate_email(email) if email else ""
    normalized_session = "" if normalized_email else (session or "").strip()
    if not normalized_email and not normalized_session:
        raise HTTPException(status_code=400, detail="Log in or continue as guest before continuing.")

    owner_field = "email" if normalized_email else "session"
    owner_value = normalized_email or normalized_session
    return normalized_email, normalized_session, owner_field, owner_value


def _owner_filter_fields(
    email: str | None = None,
    session: str | None = None,
) -> tuple[list[str], list[list[str]], str, str]:
    _, _, owner_field, owner_value = _owner_identity(email=email, session=session)
    return [owner_field], [[owner_value]], owner_field, owner_value


def _safe_uploaded_file_name(file_name: str | None) -> str:
    safe_name = Path(file_name or "").name.strip()
    if not safe_name:
        raise HTTPException(status_code=400, detail="Each uploaded document needs a file name.")
    return safe_name


def _uploaded_doc_path(uploaded_file_id: int, file_name: str) -> Path:
    suffix = Path(file_name).suffix
    return Config.DIR_CONTENTS / "uploaded_docs" / f"{uploaded_file_id}{suffix}"


def _uploaded_doc_path_by_id(uploaded_file_id: int, file_name: str | None = None) -> Path:
    expected_path = _uploaded_doc_path(uploaded_file_id, file_name or "")
    if expected_path.exists():
        return expected_path

    upload_dir = Config.DIR_CONTENTS / "uploaded_docs"
    matches = sorted(upload_dir.glob(f"{uploaded_file_id}.*"))
    if matches:
        return matches[0]
    return expected_path


def _uploaded_records_for_owner(
    email: str | None = None,
    session: str | None = None,
    limit: int = 50,
) -> list[dict[str, Any]]:
    from src import db

    _, _, owner_field, owner_value = _owner_identity(email=email, session=session)
    active_statuses = [
        status.value for status in db.uploaded_files_status if status != db.uploaded_files_status.DELETED
    ]
    df = db.selectFromDB(
        table_name="uploaded_files",
        field_names=[owner_field, "status"],
        field_values=[[owner_value], active_statuses],
        order_by_field_names=["update_date"],
        order_by_types=["DESC"],
        limit=limit,
    )
    return _records_from_dataframe(df)


def _uploaded_file_by_owner_and_name(
    file_name: str,
    email: str | None = None,
    session: str | None = None,
) -> dict[str, Any] | None:
    from src import db

    _, _, owner_field, owner_value = _owner_identity(email=email, session=session)
    df = db.selectFromDB(
        table_name="uploaded_files",
        field_names=[owner_field, "file_name", "status"],
        field_values=[[owner_value], [file_name], [db.uploaded_files_status.UPLOADED.value]],
        limit=1,
    )
    records = _records_from_dataframe(df)
    return records[0] if records else None


def _save_uploaded_file(upload: UploadFile, uploaded_file_id: int, file_name: str) -> Path:
    upload_dir = Config.DIR_CONTENTS / "uploaded_docs"
    upload_dir.mkdir(parents=True, exist_ok=True)
    destination = _uploaded_doc_path(uploaded_file_id, file_name)
    upload.file.seek(0)
    with destination.open("wb") as output_file:
        shutil.copyfileobj(upload.file, output_file)
    return destination


def _vector_collection_name(vector_db_collections_id: int) -> str:
    return f"{Config.APP_NAME_AS_PREFIX}_collection_{vector_db_collections_id}"


def _active_literature_collection_record(generated_file_id: int) -> dict[str, Any] | None:
    from src import db

    df = db.selectFromDB(
        table_name="vector_db_collections",
        field_names=["generated_files_id", "type", "status"],
        field_values=[
            [generated_file_id],
            [db.vector_db_collections_type.LITERATURE.value],
            [db.vector_db_collections_status.ACTIVE.value],
        ],
        order_by_field_names=["update_date"],
        order_by_types=["DESC"],
        limit=1,
    )
    records = _records_from_dataframe(df)
    return records[0] if records else None


def _active_literature_collection(generated_file_id: int) -> dict[str, Any] | None:
    record = _active_literature_collection_record(generated_file_id)
    if not record:
        return None
    return {
        **record,
        "collection_name": _vector_collection_name(int(record["id"])),
    }


def _active_uploaded_files_collection_record(generated_file_id: int) -> dict[str, Any] | None:
    from src import db

    df = db.selectFromDB(
        table_name="vector_db_collections",
        field_names=["generated_files_id", "type", "status"],
        field_values=[
            [generated_file_id],
            [db.vector_db_collections_type.UPLOADED_FILES.value],
            [db.vector_db_collections_status.ACTIVE.value],
        ],
        order_by_field_names=["update_date"],
        order_by_types=["DESC"],
        limit=1,
    )
    records = _records_from_dataframe(df)
    return records[0] if records else None


def _active_uploaded_files_collection(generated_file_id: int) -> dict[str, Any] | None:
    record = _active_uploaded_files_collection_record(generated_file_id)
    if not record:
        return None
    return {
        **record,
        "collection_name": _vector_collection_name(int(record["id"])),
    }


def _attached_uploaded_documents(generated_file_id: int) -> list[dict[str, Any]]:
    collection = _active_uploaded_files_collection_record(generated_file_id)
    if not collection:
        return []
    return [_uploaded_document(record) for record in _attached_uploaded_files(int(collection["id"]))]


def _uploaded_file_records_by_ids(uploaded_file_ids: list[int], email: str | None = None, session: str | None = None) -> list[dict[str, Any]]:
    from src import db

    unique_ids = sorted({int(file_id) for file_id in uploaded_file_ids})
    if not unique_ids:
        return []

    _, _, owner_field, owner_value = _owner_identity(email=email, session=session)
    df = db.selectFromDB(
        table_name="uploaded_files",
        field_names=["id", owner_field, "status"],
        field_values=[unique_ids, [owner_value], [db.uploaded_files_status.UPLOADED.value]],
    )
    return _records_from_dataframe(df)


def _uploaded_file_attachment_records(vector_db_collections_id: int) -> list[dict[str, Any]]:
    from src import db

    df = db.selectFromDB(
        table_name="vector_db_collection_files",
        field_names=["vector_db_collections_id"],
        field_values=[[vector_db_collections_id]],
    )
    return [
        record
        for record in _records_from_dataframe(df)
        if record.get("uploaded_files_id") is not None
    ]


def _attached_uploaded_files(vector_db_collections_id: int) -> list[dict[str, Any]]:
    attached_rows = _uploaded_file_attachment_records(vector_db_collections_id)
    uploaded_file_ids = [int(row["uploaded_files_id"]) for row in attached_rows if row.get("uploaded_files_id") is not None]
    if not uploaded_file_ids:
        return []

    from src import db

    df = db.selectFromDB(
        table_name="uploaded_files",
        field_names=["id"],
        field_values=[uploaded_file_ids],
    )
    records_by_id = {int(record["id"]): record for record in _records_from_dataframe(df)}
    return [records_by_id[file_id] for file_id in uploaded_file_ids if file_id in records_by_id]


def _active_uploaded_file_collections_for_uploaded_file(uploaded_file_id: int) -> list[dict[str, Any]]:
    from src import db

    df_files = db.selectFromDB(
        table_name="vector_db_collection_files",
        field_names=["uploaded_files_id"],
        field_values=[[uploaded_file_id]],
    )
    collection_ids = sorted(
        {
            int(record["vector_db_collections_id"])
            for record in _records_from_dataframe(df_files)
            if record.get("vector_db_collections_id") is not None
        }
    )
    if not collection_ids:
        return []

    df_collections = db.selectFromDB(
        table_name="vector_db_collections",
        field_names=["id", "type", "status"],
        field_values=[
            collection_ids,
            [db.vector_db_collections_type.UPLOADED_FILES.value],
            [db.vector_db_collections_status.ACTIVE.value],
        ],
        order_by_field_names=["update_date"],
        order_by_types=["DESC"],
    )
    return _records_from_dataframe(df_collections)


def _create_uploaded_files_collection(generated_file: dict[str, Any]) -> dict[str, Any]:
    from src import db
    from src.common import createVectorDBCollection

    now = datetime.now()
    generated_file_id = int(generated_file["id"])
    inserted_ids = db.insertIntoDB(
        table_name="vector_db_collections",
        field_names=[
            "email",
            "session",
            "type",
            "generated_files_id",
            "status",
            "create_date",
            "update_date",
        ],
        field_values=[
            [generated_file.get("email") or ""],
            [generated_file.get("session") or ""],
            [db.vector_db_collections_type.UPLOADED_FILES.value],
            [generated_file_id],
            [db.vector_db_collections_status.ACTIVE.value],
            [now],
            [now],
        ],
    )
    vector_db_collections_id = inserted_ids[0] if inserted_ids else None
    if vector_db_collections_id is None:
        raise HTTPException(status_code=500, detail="Unable to create an attachment collection.")

    collection_name = _vector_collection_name(int(vector_db_collections_id))
    createVectorDBCollection(collection_name=collection_name)
    return {
        "id": vector_db_collections_id,
        "email": generated_file.get("email") or "",
        "session": generated_file.get("session") or "",
        "type": db.vector_db_collections_type.UPLOADED_FILES.value,
        "generated_files_id": generated_file_id,
        "status": db.vector_db_collections_status.ACTIVE.value,
        "create_date": now,
        "update_date": now,
        "collection_name": collection_name,
    }


def _create_literature_collection(generated_file: dict[str, Any]) -> dict[str, Any]:
    from src import db
    from src.common import createVectorDBCollection

    now = datetime.now()
    generated_file_id = int(generated_file["id"])
    inserted_ids = db.insertIntoDB(
        table_name="vector_db_collections",
        field_names=[
            "email",
            "session",
            "type",
            "generated_files_id",
            "status",
            "create_date",
            "update_date",
        ],
        field_values=[
            [generated_file.get("email") or ""],
            [generated_file.get("session") or ""],
            [db.vector_db_collections_type.LITERATURE.value],
            [generated_file_id],
            [db.vector_db_collections_status.ACTIVE.value],
            [now],
            [now],
        ],
    )
    vector_db_collections_id = inserted_ids[0] if inserted_ids else None
    if vector_db_collections_id is None:
        raise HTTPException(status_code=500, detail="Unable to create the literature search collection.")

    collection_name = _vector_collection_name(int(vector_db_collections_id))
    createVectorDBCollection(collection_name=collection_name)
    return {
        "id": vector_db_collections_id,
        "email": generated_file.get("email") or "",
        "session": generated_file.get("session") or "",
        "type": db.vector_db_collections_type.LITERATURE.value,
        "generated_files_id": generated_file_id,
        "status": db.vector_db_collections_status.ACTIVE.value,
        "create_date": now,
        "update_date": now,
        "collection_name": collection_name,
    }


def _delete_vector_collection_record(collection_record: dict[str, Any], now: datetime | None = None) -> None:
    from src import db
    from src.vectordb import deleteCollection

    vector_db_collections_id = int(collection_record["id"])
    deleteCollection(_vector_collection_name(vector_db_collections_id))
    db.updateDB(
        table_name="vector_db_collections",
        update_fields=["status", "update_date"],
        update_values=[db.vector_db_collections_status.DELETED.value, now or datetime.now()],
        select_fields=["id"],
        select_values=[[vector_db_collections_id]],
    )


def _load_uploaded_files_to_vector_collection(collection_name: str, file_paths: list[tuple[str, Path]]) -> None:

    from src.vectordb import ChromaDB, getLoader

    db_vector = ChromaDB()
    try:
        db_vector.get(collection_name=collection_name)
    except Exception:
        db_vector.create(collection_name=collection_name, delete_if_exists=True)
    docs = []
    for file_name, file_path in file_paths:
        for doc in list(getLoader(file_path)):
            doc.metadata = {**{'app_file_id': Path(file_path).stem, 'app_file_type': 'uploaded_document'}, **{k: str(v) for k, v in doc.metadata.items()}}
            docs.append(doc)
    if docs:
        db_vector.add(docs=docs)


def _create_uploaded_files_collection_from_records(
    generated_file: dict[str, Any],
    uploaded_file_records: list[dict[str, Any]],
) -> dict[str, Any] | None:
    if not uploaded_file_records:
        return None

    uploaded_files_collection = _create_uploaded_files_collection(generated_file)
    vector_db_collections_id = int(uploaded_files_collection["id"])
    _insert_uploaded_file_attachment_rows(vector_db_collections_id, uploaded_file_records)
    _load_uploaded_files_to_vector_collection(
        uploaded_files_collection["collection_name"],
        _uploaded_file_paths(uploaded_file_records),
    )
    return uploaded_files_collection


def _refresh_vector_collections_for_regeneration(generated_file: dict[str, Any]) -> None:
    generated_file_id = int(generated_file["id"])
    now = datetime.now()

    uploaded_collection = _active_uploaded_files_collection_record(generated_file_id)
    if uploaded_collection:
        attached_records = _attached_uploaded_files(int(uploaded_collection["id"]))
        _delete_vector_collection_record(uploaded_collection, now)
        _create_uploaded_files_collection_from_records(generated_file, attached_records)

    literature_collection = _active_literature_collection_record(generated_file_id)
    if literature_collection:
        _delete_vector_collection_record(literature_collection, now)
        _create_literature_collection(generated_file)


def _insert_uploaded_file_attachment_rows(vector_db_collections_id: int, uploaded_file_records: list[dict[str, Any]]) -> None:
    if not uploaded_file_records:
        return

    from src import db

    existing_ids = {
        int(record["uploaded_files_id"])
        for record in _uploaded_file_attachment_records(vector_db_collections_id)
        if record.get("uploaded_files_id") is not None
    }
    new_records = [record for record in uploaded_file_records if int(record["id"]) not in existing_ids]
    if not new_records:
        return

    now = datetime.now()
    db.insertIntoDB(
        table_name="vector_db_collection_files",
        field_names=["vector_db_collections_id", "uploaded_files_id", "literature_id", "create_date", "update_date"],
        field_values=[
            [vector_db_collections_id for _ in new_records],
            [int(record["id"]) for record in new_records],
            [None for _ in new_records],
            [now for _ in new_records],
            [now for _ in new_records],
        ],
    )


def _uploaded_file_paths(records: list[dict[str, Any]]) -> list[tuple[str, Path]]:
    paths = []
    for record in records:
        file_name = str(record.get("file_name") or "")
        file_path = _uploaded_doc_path_by_id(int(record["id"]), file_name)
        if not file_path.exists():
            raise HTTPException(status_code=404, detail=f"{record.get('file_name') or 'An uploaded file'} was not found on disk.")
        paths.append((file_name, file_path))
    return paths


def _safe_download_stem(value: str | None) -> str:
    stem = re.sub(r"[^a-zA-Z0-9_-]+", "-", str(value or "generated-document")).strip("-")
    return stem or "generated-document"


def _is_content_key(value: Any) -> bool:
    return str(value).strip().lower() in {"content", "[--content--]", "__content__"}


def _text_from_content_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, (int, float, bool)):
        return str(value)
    if isinstance(value, dict):
        for key in ("body", "content", "text", "value", "generated_content"):
            if key in value:
                text = _text_from_content_value(value[key])
                if text:
                    return text
        return ""
    if isinstance(value, list):
        fragments = []
        for item in value:
            if isinstance(item, (list, tuple)) and len(item) >= 2:
                content_type = str(item[0]).strip().lower()
                if content_type in {
                    "content",
                    "body",
                    "text",
                    "generated_content",
                    "paragraph",
                    "content_user",
                    "content_ai",
                }:
                    text = _text_from_content_value(item[1])
                    if text:
                        fragments.append(text)
            elif isinstance(item, str):
                fragments.append(item.strip())
            elif isinstance(item, dict):
                text = _text_from_content_value(item)
                if text:
                    fragments.append(text)
        return "\n\n".join(fragment for fragment in fragments if fragment)
    return ""


def _section_body_from_outline_node(value: Any) -> str:
    if isinstance(value, str):
        return value.strip()
    if not isinstance(value, dict):
        return ""

    for key in ("body", "text", "generated_content"):
        if key in value:
            text = _text_from_content_value(value[key])
            if text:
                return text

    for key, item in value.items():
        if _is_content_key(key):
            return _text_from_content_value(item)

    return ""


def _heading_level(value: Any, fallback: int = 1) -> int:
    try:
        level = int(value)
    except (TypeError, ValueError):
        level = fallback
    return max(1, min(6, level))


def _manuscript_from_section_records(
    records: list[Any],
    level: int = 1,
    path: list[str] | None = None,
) -> list[dict[str, Any]]:
    sections = []
    path = path or []
    for record in records:
        if not isinstance(record, dict):
            continue

        heading = record.get("heading") or record.get("title") or record.get("header") or record.get("section")
        section_level = _heading_level(record.get("level") or record.get("depth") or record.get("heading_level"), level)
        if heading:
            current_path = [*path, str(heading)]
            body = _text_from_content_value(record.get("body") or record.get("content") or record.get("text"))
            sections.append(
                {
                    "heading": str(heading),
                    "level": section_level,
                    "body": body,
                    "raw_body": body,
                    "path": current_path,
                }
            )
        else:
            current_path = path

        children = record.get("children") or record.get("sections") or []
        if isinstance(children, list):
            sections.extend(_manuscript_from_section_records(children, section_level + 1 if heading else section_level, current_path))

    return sections


def _manuscript_from_outline_tree(
    node: Any,
    level: int = 1,
    path: list[str] | None = None,
) -> list[dict[str, Any]]:
    path = path or []
    if isinstance(node, list):
        section_records = _manuscript_from_section_records(node, level, path)
        if section_records:
            return section_records
        return []

    if not isinstance(node, dict):
        return []

    sections = []
    for heading, value in node.items():
        if _is_content_key(heading) or str(heading).strip().lower() in {"references", "bibliography"}:
            continue

        section_level = _heading_level(None, level)
        current_path = [*path, str(heading)]
        body = _section_body_from_outline_node(value)
        sections.append(
            {
                "heading": str(heading),
                "level": section_level,
                "body": body,
                "raw_body": body,
                "path": current_path,
            }
        )
        sections.extend(_manuscript_from_outline_tree(value, section_level + 1, current_path))

    return sections


def _manuscript_from_outline_file(file_id: int) -> tuple[list[dict[str, Any]], bool]:
    outline_file_path = _outline_file_path(file_id)
    if not outline_file_path.exists():
        return [], False

    outline_data = _load_outline_json(outline_file_path)

    if isinstance(outline_data, dict):
        if any(key in outline_data for key in ("heading", "title", "header", "section")):
            return _manuscript_from_section_records([outline_data]), True
        for key in ("manuscript", "sections"):
            if isinstance(outline_data.get(key), list):
                return _manuscript_from_section_records(outline_data[key]), True

    return _manuscript_from_outline_tree(outline_data), True


def _raw_outline_from_outline_file(file_id: int) -> tuple[str, bool]:
    outline_file_path = Config.DIR_CONTENTS / f"outline_{file_id}.json"
    if not outline_file_path.exists():
        return "", False

    outline_data = _load_outline_json(outline_file_path)

    return _raw_outline_from_outline_data(outline_data), True


def _outline_template_content(template_name: str) -> str:
    if not re.match(r"^[a-zA-Z0-9_-]+$", template_name):
        raise HTTPException(status_code=400, detail="Invalid outline template name.")

    template_path = _outline_templates_dir() / f"{template_name}.md"
    if not template_path.exists():
        raise HTTPException(status_code=404, detail="Outline template was not found.")

    return template_path.read_text()


def _outline_templates_dir() -> Path:
    return Path(__file__).resolve().parent / "data" / "outline_templates"


def _outline_templates() -> list[dict[str, str]]:
    
    def _get_template_name(str):
        str = str.replace('_', ' ')
        return str[0].upper() + str[1:]
    
    template_dir = _outline_templates_dir()
    if not template_dir.exists():
        return []

    return [
        {
            "name": template_path.stem,
            "label": _get_template_name(template_path.stem),
        }
        for template_path in sorted(template_dir.glob("*.md"))
        if re.match(r"^[a-zA-Z0-9_-]+$", template_path.stem)
    ]


def _validated_credentials_id(credentials_id: int | None) -> int:
    if credentials_id is None or credentials_id <= 0:
        raise HTTPException(status_code=400, detail="Start a workspace session before using uploaded templates.")
    if _credential_by_id(credentials_id) is None:
        raise HTTPException(status_code=404, detail="The account for these uploaded templates was not found.")
    return int(credentials_id)


def _user_outline_templates_dir(credentials_id: int) -> Path:
    return Config.DIR_CONTENTS / f"user_{credentials_id}" / "outline_templates"


def _uploaded_outline_template_file_name(file_name: str | None) -> str:
    safe_name = _safe_uploaded_file_name(file_name)
    template_path = Path(safe_name)
    if template_path.suffix.lower() not in OUTLINE_IMPORT_ALLOWED_SUFFIXES:
        raise HTTPException(status_code=400, detail="Outline templates must be Markdown (.md) or Word (.docx) files.")

    stem = re.sub(r"[^a-zA-Z0-9_-]+", "_", template_path.stem).strip("_")
    if not stem:
        raise HTTPException(status_code=400, detail="Each uploaded template needs a valid file name.")
    return f"{stem}.md"


def _uploaded_outline_templates(credentials_id: int) -> list[dict[str, Any]]:
    template_dir = _user_outline_templates_dir(credentials_id)
    if not template_dir.exists():
        return []

    return [
        {
            "name": template_path.stem,
            "label": template_path.stem.replace("_", " ").replace("-", " ").title(),
            "file_name": template_path.name,
            "last_modified": datetime.fromtimestamp(template_path.stat().st_mtime).isoformat(),
        }
        for template_path in sorted(template_dir.glob("*.md"))
        if re.match(r"^[a-zA-Z0-9_-]+$", template_path.stem)
    ]


def _uploaded_outline_template_content(credentials_id: int, template_name: str) -> str:
    if not re.match(r"^[a-zA-Z0-9_-]+$", template_name):
        raise HTTPException(status_code=400, detail="Invalid uploaded outline template name.")

    template_path = _user_outline_templates_dir(credentials_id) / f"{template_name}.md"
    if not template_path.exists():
        raise HTTPException(status_code=404, detail="Uploaded outline template was not found.")

    return template_path.read_text(encoding="utf-8")


async def _save_uploaded_outline_template(credentials_id: int, upload: UploadFile) -> dict[str, Any]:
    file_name = _uploaded_outline_template_file_name(upload.filename)
    content = await upload.read()
    if not content:
        raise HTTPException(status_code=400, detail=f"{file_name} is empty.")

    outline = _outline_from_uploaded_file(upload.filename, content)
    template_dir = _user_outline_templates_dir(credentials_id)
    template_dir.mkdir(parents=True, exist_ok=True)
    template_path = template_dir / file_name
    template_path.write_text(outline, encoding="utf-8")
    return {
        "name": template_path.stem,
        "label": template_path.stem.replace("_", " ").replace("-", " ").title(),
        "file_name": template_path.name,
        "source_file_name": upload.filename,
        "content": outline,
        "last_modified": datetime.fromtimestamp(template_path.stat().st_mtime).isoformat(),
    }


def _uploaded_outline_template_path(credentials_id: int, template_name: str) -> Path:
    if not re.match(r"^[a-zA-Z0-9_-]+$", template_name):
        raise HTTPException(status_code=400, detail="Invalid uploaded outline template name.")
    return _user_outline_templates_dir(credentials_id) / f"{template_name}.md"


def _write_uploaded_outline_template_content(credentials_id: int, template_name: str, outline: str) -> dict[str, Any]:
    outline_text = str(outline or "").strip()
    if not outline_text:
        raise HTTPException(status_code=400, detail="Outline template content cannot be empty.")

    from src.manage_outline import getRawOutline, processOutline

    try:
        normalized_outline = "\n".join(getRawOutline(processOutline(outline_text), raw_outline=[])).strip()
    except Exception:
        logger.exception("Unable to validate uploaded outline template %s", template_name)
        raise HTTPException(
            status_code=400,
            detail="We could not save that outline. Check the heading levels and try again.",
        )

    template_path = _uploaded_outline_template_path(credentials_id, template_name)
    if not template_path.exists():
        raise HTTPException(status_code=404, detail="Uploaded outline template was not found.")
    template_path.write_text(normalized_outline, encoding="utf-8")
    return {
        "name": template_path.stem,
        "label": template_path.stem.replace("_", " ").replace("-", " ").title(),
        "file_name": template_path.name,
        "last_modified": datetime.fromtimestamp(template_path.stat().st_mtime).isoformat(),
    }


def _delete_uploaded_outline_template(credentials_id: int, template_name: str) -> dict[str, Any]:
    template_path = _uploaded_outline_template_path(credentials_id, template_name)
    if not template_path.exists():
        raise HTTPException(status_code=404, detail="Uploaded outline template was not found.")
    template = {
        "name": template_path.stem,
        "label": template_path.stem.replace("_", " ").replace("-", " ").title(),
        "file_name": template_path.name,
    }
    template_path.unlink()
    return template


def _rename_uploaded_outline_template(credentials_id: int, template_name: str, file_name: str) -> dict[str, Any]:
    current_path = _uploaded_outline_template_path(credentials_id, template_name)
    if not current_path.exists():
        raise HTTPException(status_code=404, detail="Uploaded outline template was not found.")

    next_file_name = _uploaded_outline_template_file_name(file_name)
    next_path = _user_outline_templates_dir(credentials_id) / next_file_name
    if current_path.resolve() != next_path.resolve() and next_path.exists():
        raise HTTPException(status_code=409, detail=f"An uploaded outline template named {next_file_name} already exists.")

    current_path.rename(next_path)
    return {
        "name": next_path.stem,
        "label": next_path.stem.replace("_", " ").replace("-", " ").title(),
        "file_name": next_path.name,
        "last_modified": datetime.fromtimestamp(next_path.stat().st_mtime).isoformat(),
    }


def _process_outline(outline: str) -> dict[str, Any]:
    from src.manage_outline import processOutline

    return processOutline(outline)


def _raw_outline_from_outline_data(outline_data: dict[str, Any]) -> str:
    from src.manage_outline import getRawOutline

    raw_outline_lines = getRawOutline(outline_data, [], 1)
    return "\n".join(str(line) for line in raw_outline_lines).strip()


GENERATION_JOBS: dict[str, dict[str, Any]] = {}
DOC_CONTENT_LOCK = threading.Lock()
OUTLINE_CONTENT_KEY = "content"
OUTLINE_IS_ABSTRACT = "is_abstract"
OUTLINE_INSTRUCTIONS = "instructions"
OUTLINE_CONTENT_USER = "content_user"
OUTLINE_CONTENT_AI = "content_ai"
OUTLINE_CONTENT_PRE_SUMMARY = "content_pre_summary"
OUTLINE_CONCEPT_MAP = "concept_map"


def _outline_file_path(generated_file_id: int) -> Path:
    return Config.DIR_CONTENTS / f"outline_{generated_file_id}.json"


def _load_outline_json(outline_file_path: Path) -> dict[str, Any]:
    with outline_file_path.open() as fp:
        try:
            outline_data = json.load(fp)
        except json.JSONDecodeError as exp:
            logger.exception("Saved outline JSON is invalid: %s", outline_file_path)
            raise HTTPException(
                status_code=409,
                detail="The saved outline file is empty or invalid. Save the outline again, then reload this manuscript.",
            ) from exp

    if not isinstance(outline_data, dict):
        raise HTTPException(
            status_code=409,
            detail="The saved outline file is not in the expected format. Save the outline again, then reload this manuscript.",
        )
    return outline_data


def _read_outline_file(generated_file_id: int) -> dict[str, Any]:
    return _load_outline_json(_outline_file_path(generated_file_id))


def _write_outline_file(generated_file_id: int, outline_data: dict[str, Any]) -> Path:
    outline_file_path = _outline_file_path(generated_file_id)
    outline_file_path.parent.mkdir(parents=True, exist_ok=True)
    temp_outline_file_path = outline_file_path.with_name(f"{outline_file_path.name}.{uuid4().hex}.tmp")
    try:
        with temp_outline_file_path.open("w") as fp:
            json.dump(outline_data, fp, indent=2)
            fp.write("\n")
        temp_outline_file_path.replace(outline_file_path)
    finally:
        if temp_outline_file_path.exists():
            temp_outline_file_path.unlink()
    return outline_file_path


def _save_processed_outline(
    generated_file_id: int,
    outline: str,
    email: str | None = None,
    session: str | None = None,
    preserve_generated_content: bool = False,
) -> tuple[dict[str, Any], dict[str, Any], Path]:
    from src import db

    outline = outline.strip()
    if not outline:
        raise HTTPException(status_code=400, detail="A structured outline is required before content generation can start.")

    record = _generated_file_by_id(generated_file_id, email=email, session=session)
    try:
        processed_outline = _process_outline(outline)
    except Exception as exp:
        logger.exception("Unable to parse structured outline")
        raise HTTPException(
            status_code=400,
            detail="I could not understand that structured outline. Check the heading levels and [--content--] tags, then try again.",
        ) from exp

    if preserve_generated_content:
        try:
            existing_outline = _read_outline_file(generated_file_id)
            _merge_existing_generated_content(processed_outline, existing_outline)
        except FileNotFoundError:
            pass

    outline_file_path = _write_outline_file(generated_file_id, processed_outline)
    now = datetime.now()
    db.updateDB(
        table_name="generated_files",
        update_fields=["update_date"],
        update_values=[now],
        select_fields=["id"],
        select_values=[[generated_file_id]],
    )
    return {**record, "update_date": now}, processed_outline, outline_file_path


def _set_generated_file_status(generated_file_id: int, status: str) -> None:
    try:
        from src import db

        db.updateDB(
            table_name="generated_files",
            update_fields=["status", "update_date"],
            update_values=[status, datetime.now()],
            select_fields=["id"],
            select_values=[[generated_file_id]],
        )
    except Exception:
        logger.exception("Unable to update generated file %s status to %s", generated_file_id, status)


def _resolve_generation_architecture(
    generated_file_id: int,
    architecture_type: str,
    collection_name: str = "",
    collection_name_lit_search: str = "",
) -> tuple[str, str, str]:
    from src import db

    if architecture_type != db.generated_files_ai_architecture.RAG.value:
        return architecture_type, collection_name, collection_name_lit_search

    uploaded_collection = _active_uploaded_files_collection(generated_file_id)
    literature_collection = _active_literature_collection(generated_file_id)
    collection_name = uploaded_collection["collection_name"] if uploaded_collection else collection_name
    collection_name_lit_search = (
        literature_collection["collection_name"] if literature_collection else collection_name_lit_search
    )

    if not collection_name and not collection_name_lit_search:
        logger.warning(
            "Generated file %s is marked as RAG but has no active uploaded-file or literature-search collection. Falling back to base generation.",
            generated_file_id,
        )
        db.updateDB(
            table_name="generated_files",
            update_fields=["ai_architecture", "update_date"],
            update_values=[db.generated_files_ai_architecture.BASE.value, datetime.now()],
            select_fields=["id"],
            select_values=[[generated_file_id]],
        )
        return db.generated_files_ai_architecture.BASE.value, "", ""

    return architecture_type, collection_name, collection_name_lit_search


def _job_snapshot(job: dict[str, Any]) -> dict[str, Any]:
    public_keys = {
        "id",
        "generated_file_id",
        "status",
        "message",
        "error",
        "current_section",
        "completed_sections",
        "total_sections",
        "manuscript",
        "concept_maps",
        "attached_reference_list_from_content",
        "generated_file",
        "created_at",
        "updated_at",
    }
    return _jsonable({key: job.get(key) for key in public_keys if key in job})


def _active_generation_job_for_file(generated_file_id: int) -> dict[str, Any] | None:
    for job in GENERATION_JOBS.values():
        if job.get("generated_file_id") == generated_file_id and job.get("worker_active"):
            return job
    return None


def _finalize_generation_task(job: dict[str, Any]) -> None:
    job["task"] = None
    if job.get("pause_requested") and job.get("status") not in {"completed", "error"}:
        job["status"] = "paused"
        job["message"] = "Generation paused. Click Generate to continue with the remaining outline."
        job["current_section"] = ""
    if job.get("status") in {"completed", "error", "paused"}:
        job["worker_active"] = False
    job["updated_at"] = datetime.now().isoformat()


def _normalize_content_item(item: Any) -> tuple[str, Any] | None:
    if isinstance(item, (list, tuple)) and len(item) >= 2:
        return str(item[0]), item[1]
    return None


def _content_items(node: dict[str, Any]) -> list[Any]:
    content = node.get(OUTLINE_CONTENT_KEY)
    return content if isinstance(content, list) else []


def _content_value(items: list[Any], content_type: str, default: Any = "") -> Any:
    for item in items:
        normalized = _normalize_content_item(item)
        if normalized and normalized[0] == content_type:
            return normalized[1]
    return default


def _normalize_concept_map(value: Any) -> dict[str, list[str]]:
    if not isinstance(value, dict):
        return {}

    concept_map: dict[str, list[str]] = {}
    for key, raw_children in value.items():
        node_label = str(key).strip()
        if not node_label:
            continue

        if isinstance(raw_children, (list, tuple, set)):
            children = [str(child).strip() for child in raw_children if str(child).strip()]
        elif raw_children in (None, "", {}):
            children = []
        else:
            children = [str(raw_children).strip()] if str(raw_children).strip() else []

        concept_map[node_label] = children

    return concept_map


def _concept_maps_from_outline_tree(node: Any, path: list[str] | None = None) -> list[dict[str, Any]]:
    if not isinstance(node, dict):
        return []

    path = path or []
    concept_maps: list[dict[str, Any]] = []
    items = _content_items(node)
    concept_map = _normalize_concept_map(_content_value(items, OUTLINE_CONCEPT_MAP, {}))
    if concept_map:
        concept_maps.append(
            {
                "section": path[-1] if path else "Concept map",
                "path": path.copy(),
                "map": concept_map,
            }
        )

    for key, value in node.items():
        if key == OUTLINE_CONTENT_KEY:
            continue
        concept_maps.extend(_concept_maps_from_outline_tree(value, [*path, str(key)]))

    return concept_maps


def _concept_maps_from_outline_file(file_id: int) -> tuple[list[dict[str, Any]], bool]:
    outline_file_path = _outline_file_path(file_id)
    if not outline_file_path.exists():
        return [], False

    outline_data = _load_outline_json(outline_file_path)

    return _concept_maps_from_outline_tree(outline_data), True


def _normalize_attached_references_content(value: Any) -> list[str]:
    if value in (None, "", {}):
        return []
    if isinstance(value, str):
        return [value.strip()] if value.strip() else []
    if not isinstance(value, list):
        return [str(value).strip()] if str(value).strip() else []

    references = []
    seen = set()
    for item in value:
        text = str(item or "").strip()
        if not text or text in seen:
            continue
        seen.add(text)
        references.append(text)
    return references


def _attached_reference_map_for_generated_file(generated_file_id: int) -> dict[str, str]:
    from src.common import getAttachedRefs

    uploaded_collection = _active_uploaded_files_collection_record(generated_file_id)
    literature_collection = _active_literature_collection_record(generated_file_id)
    uploaded_collection_id = int(uploaded_collection["id"]) if uploaded_collection else None
    literature_collection_id = int(literature_collection["id"]) if literature_collection else None
    attached_references_db, _ = getAttachedRefs(uploaded_collection_id, literature_collection_id)
    return {str(reference_id): reference_text for reference_id, reference_text, _reference_type in attached_references_db}


def _process_manuscript_citations(
    manuscript: list[dict[str, Any]],
    attached_references_db: dict[str, str],
) -> tuple[list[dict[str, Any]], list[str]]:
    from src.common import processCitation, sanitizeContent

    current_attached_references_content = []
    processed_manuscript = []
    for section in manuscript:
        section_body = str(section.get("body") or "")
        if attached_references_db and "CITE(" in section_body:
            section_body, current_attached_references_content = processCitation(section_body, attached_references_db, current_attached_references_content, enable_html_link_format=True)
        section_body = sanitizeContent(section_body)
        processed_manuscript.append({**section, "body": section_body})

    return processed_manuscript, current_attached_references_content


def _truthy_content_flag(value: Any) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, str):
        return value.strip().lower() in {"1", "true", "yes", "abstract"}
    return bool(value)


def _set_content_value(items: list[Any], content_type: str, value: Any) -> None:
    for item in items:
        if isinstance(item, list) and item and str(item[0]) == content_type:
            if len(item) == 1:
                item.append(value)
            else:
                item[1] = value
            return
        if isinstance(item, tuple) and len(item) >= 2 and str(item[0]) == content_type:
            items[items.index(item)] = [content_type, value]
            return
    items.append([content_type, value])


def _outline_node_for_path(outline_data: dict[str, Any], path: list[str] | tuple[str, ...]) -> dict[str, Any] | None:
    node: Any = outline_data
    for heading in path:
        if not isinstance(node, dict):
            return None
        node = node.get(heading)
    return node if isinstance(node, dict) else None


def _outline_node_for_heading_and_paragraph(
    outline_data: dict[str, Any],
    heading: str,
    raw_paragraph: str,
) -> dict[str, Any] | None:
    heading = str(heading or "").strip()
    raw_paragraph = str(raw_paragraph or "").strip()
    if not heading:
        return None

    matches: list[dict[str, Any]] = []

    def walk(node: Any) -> None:
        if not isinstance(node, dict):
            return

        for key, value in node.items():
            if key == OUTLINE_CONTENT_KEY:
                continue
            if not isinstance(value, dict):
                continue

            if str(key).strip() == heading:
                items = _content_items(value)
                raw_content = str(_content_value(items, OUTLINE_CONTENT_AI, "") or "")
                if not raw_paragraph or raw_paragraph in raw_content or not raw_content:
                    matches.append(value)
            walk(value)

    walk(outline_data)
    if len(matches) == 1:
        return matches[0]
    return None


def _display_manuscript_from_outline_tree(
    outline_data: dict[str, Any],
    attached_references_db: dict[str, str],
    content_overrides: dict[tuple[str, ...], Any] | None = None,
) -> tuple[list[dict[str, Any]], list[str]]:
    display_outline = deepcopy(outline_data)
    for path, content in (content_overrides or {}).items():
        node = _outline_node_for_path(display_outline, path)
        if node is None:
            continue
        items = _content_items(node)
        if not items:
            node[OUTLINE_CONTENT_KEY] = []
            items = node[OUTLINE_CONTENT_KEY]
        _set_content_value(items, OUTLINE_CONTENT_AI, content)

    manuscript = _manuscript_from_outline_tree(display_outline)
    return _process_manuscript_citations(manuscript, attached_references_db)


def _has_content_value(items: list[Any], content_type: str) -> bool:
    value = _content_value(items, content_type, "")
    if isinstance(value, str):
        return bool(value.strip())
    return value not in (None, "", [], {})


def _section_has_generated_content(items: list[Any]) -> bool:
    return _has_content_value(items, OUTLINE_CONTENT_AI)


def _merge_existing_generated_content(target_node: Any, source_node: Any) -> None:
    if not isinstance(target_node, dict) or not isinstance(source_node, dict):
        return

    target_items = _content_items(target_node)
    source_items = _content_items(source_node)
    if target_items and source_items:
        for content_type in (OUTLINE_IS_ABSTRACT, OUTLINE_CONTENT_AI, OUTLINE_CONTENT_PRE_SUMMARY, OUTLINE_CONCEPT_MAP):
            if _has_content_value(source_items, content_type):
                _set_content_value(target_items, content_type, _content_value(source_items, content_type))

    for key, target_value in target_node.items():
        if key == OUTLINE_CONTENT_KEY:
            continue
        _merge_existing_generated_content(target_value, source_node.get(key))


def _clear_generated_content_items(items: list[Any]) -> bool:
    changed = False
    for index, item in enumerate(items):
        normalized = _normalize_content_item(item)
        if not normalized:
            continue

        content_type, value = normalized
        if content_type not in {OUTLINE_CONTENT_AI, OUTLINE_CONTENT_PRE_SUMMARY, OUTLINE_CONCEPT_MAP}:
            continue

        empty_value: Any = {} if content_type == OUTLINE_CONCEPT_MAP else ""
        if value == empty_value:
            continue

        if isinstance(item, list):
            if len(item) == 1:
                item.append(empty_value)
            else:
                item[1] = empty_value
        else:
            items[index] = [content_type, empty_value]
        changed = True

    return changed


def _reset_outline_generated_content(node: Any) -> bool:
    changed = False
    if isinstance(node, dict):
        items = _content_items(node)
        if items:
            changed = _clear_generated_content_items(items) or changed

        for key, value in node.items():
            if key == OUTLINE_CONTENT_KEY:
                continue
            changed = _reset_outline_generated_content(value) or changed
    elif isinstance(node, list):
        for item in node:
            changed = _reset_outline_generated_content(item) or changed

    return changed


def _reset_generated_document_content(generated_file_id: int) -> tuple[list[dict[str, Any]], list[dict[str, Any]], list[str], bool]:
    try:
        outline_data = _read_outline_file(generated_file_id)
    except FileNotFoundError:
        return [], [], [], False

    if _reset_outline_generated_content(outline_data):
        _write_outline_file(generated_file_id, outline_data)

    return _manuscript_from_outline_tree(outline_data), [], [], True


def _section_needs_ai(items: list[Any]) -> bool:
    return any((_normalize_content_item(item) or ("", ""))[0] == OUTLINE_CONTENT_AI for item in items)


def _section_is_abstract(items: list[Any]) -> bool:
    return _truthy_content_flag(_content_value(items, OUTLINE_IS_ABSTRACT, False))


def _heading_looks_like_abstract(heading: str) -> bool:
    normalized = re.sub(r"[^a-z]+", " ", heading.lower()).strip()
    return normalized in {"abstract", "summary", "executive summary"} or normalized.startswith("abstract ")


def _detector_response_is_abstract(response: Any) -> bool:
    if isinstance(response, dict):
        return _truthy_content_flag(response.get(OUTLINE_IS_ABSTRACT, False))
    return _truthy_content_flag(getattr(response, OUTLINE_IS_ABSTRACT, False))


def _section_instructions(items: list[Any]) -> str:
    return str(_content_value(items, OUTLINE_INSTRUCTIONS, "") or "").strip()


def _section_prompt(path: list[str], items: list[Any]) -> str:
    lines = []
    for index, heading in enumerate(path):
        lines.append(f"{'#' * min(index + 1, 6)} {heading}")

    for item in items:
        normalized = _normalize_content_item(item)
        if not normalized:
            continue
        content_type, text = normalized
        if content_type == OUTLINE_CONTENT_USER and text:
            lines.append(str(text).strip())
        elif content_type == OUTLINE_CONTENT_AI:
            lines.append("[--content--]")

    return "\n\n".join(line for line in lines if line)


def _generation_section_record(path: list[str], node: dict[str, Any], items: list[Any]) -> dict[str, Any]:
    return {
        "path": path.copy(),
        "heading": path[-1] if path else "Untitled section",
        "node": node,
        "items": items,
        "instructions": _section_instructions(items),
        "is_abstract": _section_is_abstract(items),
        "current_section": _section_prompt(path, items),
    }


def _generation_sections_in_outline_order(outline_data: dict[str, Any]) -> list[dict[str, Any]]:
    sections: list[dict[str, Any]] = []

    def walk(node: Any, path: list[str]) -> None:
        if not isinstance(node, dict):
            return

        items = _content_items(node)
        if items and _section_needs_ai(items):
            sections.append(_generation_section_record(path, node, items))

        for key, value in node.items():
            if key == OUTLINE_CONTENT_KEY:
                continue
            walk(value, [*path, str(key)])

    walk(outline_data, [])
    return sections


def _content_summary_text(value: Any) -> str:
    if isinstance(value, str):
        return value.strip()
    return _text_from_content_value(value).strip()


def _last_completed_regular_section_summary(outline_data: dict[str, Any]) -> str:
    last_section: dict[str, Any] | None = None
    for section in _generation_sections_in_outline_order(outline_data):
        if section["is_abstract"] or not _section_has_generated_content(section["items"]):
            continue
        last_section = section

    if not last_section:
        return ""

    summary = _content_summary_text(_content_value(last_section["items"], OUTLINE_CONTENT_PRE_SUMMARY, ""))
    if summary:
        return summary

    return _content_summary_text(_content_value(last_section["items"], OUTLINE_CONTENT_AI, ""))


def _summary_before_section(outline_data: dict[str, Any], target_path: list[str]) -> str:
    target_path_tuple = tuple(target_path)
    previous_summary = ""
    for section in _generation_sections_in_outline_order(outline_data):
        if tuple(section["path"]) == target_path_tuple:
            return previous_summary
        if section["is_abstract"] or not _section_has_generated_content(section["items"]):
            continue
        summary = _content_summary_text(_content_value(section["items"], OUTLINE_CONTENT_PRE_SUMMARY, ""))
        previous_summary = summary or _content_summary_text(_content_value(section["items"], OUTLINE_CONTENT_AI, ""))
    return previous_summary


def _initial_content_pre_summary(outline_data: dict[str, Any], sections: list[dict[str, Any]], mode: str) -> str:
    if mode != "remaining" or not sections:
        return ""
    if sections[0]["is_abstract"]:
        return _last_completed_regular_section_summary(outline_data)
    return _summary_before_section(outline_data, sections[0]["path"])


def _split_manuscript_paragraphs(text: Any) -> list[str]:
    normalized = str(text or "").replace("\r\n", "\n").strip()
    if not normalized:
        return []

    blocks = [block.strip() for block in re.split(r"\n\s*\n+", normalized) if block.strip()]
    if len(blocks) > 1:
        return blocks

    return [block.strip() for block in normalized.split("\n") if block.strip()]


def _paragraph_update_instruction(action: str) -> str:
    if action == "Expand":
        return (
            "Write additional scholarly text that expands on the selected paragraph. "
            "Do not repeat the selected paragraph. Return only the new text to append after it."
        )
    if action == "Rephrase":
        return (
            "Rephrase the selected paragraph for clarity, flow, and scholarly tone while preserving the meaning. "
            "Return only the replacement paragraph text."
        )
    return "Return only the replacement paragraph text."


def _paragraph_update_prompt(path: list[str], section_body: str, paragraph_index: int, action: str) -> str:
    paragraphs = _split_manuscript_paragraphs(section_body)
    if paragraph_index >= len(paragraphs):
        raise HTTPException(status_code=400, detail="The selected paragraph could not be found in the saved manuscript.")
    if action == "Expand":
        paragraphs.insert(paragraph_index + 1, "[--content--]")
    else:
        paragraphs[paragraph_index] = "[--content--]"
    lines = [f"{'#' * min(index + 1, 6)} {heading}" for index, heading in enumerate(path)]
    lines.append("\n\n".join(paragraphs))
    return "\n\n".join(line for line in lines if str(line).strip())


def _content_response_text(response: Any) -> str:
    if isinstance(response, dict):
        content = response.get("content")
        if content is None and isinstance(response.get("result"), dict):
            content = response["result"].get("content")
        return _text_from_content_value(content).strip()
    return _text_from_content_value(response).strip()


def _replace_paragraph_in_section_body(
    section_body: str,
    paragraph_index: int,
    raw_paragraph: str,
    replacement: str,
    action: str,
) -> str:
    paragraphs = _split_manuscript_paragraphs(section_body)
    sent_paragraph = str(raw_paragraph or "").strip()
    if sent_paragraph and (paragraph_index >= len(paragraphs) or paragraphs[paragraph_index].strip() != sent_paragraph):
        try:
            paragraph_index = next(
                index
                for index, paragraph in enumerate(paragraphs)
                if paragraph.strip() == sent_paragraph
            )
        except StopIteration:
            raise HTTPException(
                status_code=409,
                detail="The selected paragraph has changed. Reload the manuscript and try again.",
            )

    if paragraph_index >= len(paragraphs):
        raise HTTPException(status_code=400, detail="The selected paragraph could not be found in the saved manuscript.")

    replacement = str(replacement or "").strip()
    if action == "Expand" and replacement:
        if sent_paragraph and replacement.startswith(sent_paragraph):
            replacement = replacement[len(sent_paragraph):].strip()
        if replacement:
            paragraphs.insert(paragraph_index + 1, replacement)
    elif replacement:
        paragraphs[paragraph_index] = replacement
    else:
        paragraphs.pop(paragraph_index)

    return "\n\n".join(paragraph for paragraph in paragraphs if paragraph.strip())


def _extract_generation_sections(outline_data: dict[str, Any], remaining_only: bool = True) -> list[dict[str, Any]]:
    sections = [
        section
        for section in _generation_sections_in_outline_order(outline_data)
        if not remaining_only or not _section_has_generated_content(section["items"])
    ]
    regular_sections = [section for section in sections if not section["is_abstract"]]
    abstract_sections = [section for section in sections if section["is_abstract"]]
    return regular_sections + abstract_sections


def _mark_first_abstract_section(outline_data: dict[str, Any], agent_abstract_detector: Any) -> tuple[dict[str, Any], str]:
    if not outline_data:
        return outline_data, ""

    title_node = next(iter(outline_data.values()), None)
    if not isinstance(title_node, dict):
        return outline_data, ""

    first_section_heading = ""
    first_section_node: dict[str, Any] | None = None
    for heading, node in title_node.items():
        if heading == OUTLINE_CONTENT_KEY:
            continue
        if isinstance(node, dict):
            first_section_heading = str(heading)
            first_section_node = node
            break

    if not first_section_heading or first_section_node is None:
        return outline_data, ""

    items = _content_items(first_section_node)
    if items and _section_is_abstract(items):
        return outline_data, first_section_heading

    is_abstract = _heading_looks_like_abstract(first_section_heading)
    if not is_abstract and _section_needs_ai(items):
        response = agent_abstract_detector.invoke({"current_section": first_section_heading})
        is_abstract = _detector_response_is_abstract(response)

    if not is_abstract:
        return outline_data, ""

    if not isinstance(first_section_node.get(OUTLINE_CONTENT_KEY), list):
        first_section_node[OUTLINE_CONTENT_KEY] = []
    items = first_section_node[OUTLINE_CONTENT_KEY]
    if not _section_needs_ai(items):
        _set_content_value(items, OUTLINE_CONTENT_AI, "")
    if not _section_is_abstract(items):
        items.insert(0, [OUTLINE_IS_ABSTRACT, True])
    return outline_data, first_section_heading


def _safe_architecture_classes() -> tuple[Any, Any, Any]:
    _required_default_model()
    original_apply = None
    nest_asyncio_module = None
    try:
        import nest_asyncio

        nest_asyncio_module = nest_asyncio
        original_apply = nest_asyncio.apply
        nest_asyncio.apply = lambda *args, **kwargs: None
    except Exception:
        nest_asyncio_module = None

    try:
        from src.ai.architecture import (
            AbstractSectionDetectorArchitecture,
            AbstractWriterArchitecture,
            ContentWriterArchitecture,
        )

        return AbstractSectionDetectorArchitecture, AbstractWriterArchitecture, ContentWriterArchitecture
    finally:
        if nest_asyncio_module is not None and original_apply is not None:
            nest_asyncio_module.apply = original_apply


async def _run_generation_job(
    job_id: str,
    generated_file_id: int,
    request: GeneratedFileGenerateRequest,
    generated_file: dict[str, Any],
) -> None:
    job = GENERATION_JOBS[job_id]
    outline_data: dict[str, Any] | None = None
    attached_references_db = request.attached_references_db.copy()
    display_content_overrides: dict[tuple[str, ...], Any] = {}

    def update_job(**updates: Any) -> None:
        job.update(updates)
        job["updated_at"] = datetime.now().isoformat()

    def manuscript_snapshot(
        outline_data_current: dict[str, Any],
        attached_references_content_override: Any = None,
    ) -> tuple[list[dict[str, Any]], list[str]]:
        manuscript, processed_attached_references_content = _display_manuscript_from_outline_tree(
            outline_data_current,
            attached_references_db,
            display_content_overrides,
        )
        return manuscript, _normalize_attached_references_content(attached_references_content_override) or processed_attached_references_content

    try:
        from src import db
        from src.generate import generateContent

        def pause_if_requested(outline_data_current: dict[str, Any], completed_sections: int) -> bool:
            if not job.get("pause_requested"):
                return False
            manuscript, attached_reference_list_from_content = manuscript_snapshot(outline_data_current, job.get("attached_reference_list_from_content", []))
            _set_generated_file_status(generated_file_id, db.generated_files_status.CANCELLED.value)
            update_job(
                status="paused",
                message="Generation paused. Click Generate to continue with the remaining outline.",
                current_section="",
                completed_sections=completed_sections,
                worker_active=False,
                generated_file={**generated_file, "status": db.generated_files_status.CANCELLED.value},
                manuscript=manuscript,
                attached_reference_list_from_content=attached_reference_list_from_content,
            )
            return True

        update_job(status="running", message="Reading the saved outline...")
        _set_generated_file_status(generated_file_id, db.generated_files_status.RUNNING.value)
        outline_data = _read_outline_file(generated_file_id)
        if pause_if_requested(outline_data, 0):
            return

        update_job(message="Checking for an abstract section...")
        AbstractSectionDetectorArchitecture, AbstractWriterArchitecture, ContentWriterArchitecture = _safe_architecture_classes()
        detector = AbstractSectionDetectorArchitecture(
            model_name=_model_name(request.model_name),
            temperature=request.temperature,
            instructions=request.instructions,
        )
        outline_data, _ = await asyncio.to_thread(_mark_first_abstract_section, outline_data, detector)
        _write_outline_file(generated_file_id, outline_data)

        sections = _extract_generation_sections(outline_data, remaining_only=request.mode == "remaining")
        manuscript, attached_reference_list_from_content = manuscript_snapshot(outline_data)
        update_job(
            total_sections=len(sections),
            manuscript=manuscript,
            attached_reference_list_from_content=attached_reference_list_from_content,
        )
        if not sections:
            _set_generated_file_status(generated_file_id, db.generated_files_status.SUCCESS.value)
            update_job(
                status="completed",
                message="The structured outline was saved. No AI content sections were marked for generation.",
                current_section="",
                worker_active=False,
                generated_file={**generated_file, "status": db.generated_files_status.SUCCESS.value},
                manuscript=manuscript,
                attached_reference_list_from_content=attached_reference_list_from_content,
            )
            return

        architecture_type = generated_file.get("ai_architecture") or request.architecture_type or "base"
        architecture_type, collection_name, collection_name_lit_search = _resolve_generation_architecture(
            generated_file_id=generated_file_id,
            architecture_type=architecture_type,
            collection_name=request.collection_name,
            collection_name_lit_search=request.collection_name_lit_search,
        )
        generated_file = {
            **generated_file,
            "ai_architecture": architecture_type,
        }

        writer = ContentWriterArchitecture(
            model_name=_model_name(request.model_name),
            temperature=request.temperature,
            instructions=request.instructions,
            type=architecture_type,
            collection_name=collection_name,
            collection_name_lit_search=collection_name_lit_search,
        )
        abstract_writer = AbstractWriterArchitecture(
            model_name=_model_name(request.model_name),
            temperature=request.temperature,
            instructions=request.instructions,
        )

        content_pre_summary = _initial_content_pre_summary(outline_data, sections, request.mode)
        attached_references_db = _attached_reference_map_for_generated_file(generated_file_id)
        section_attached_references_content = []
        for index, section in enumerate(sections, start=1):
            if pause_if_requested(outline_data, index - 1):
                return

            section_label = section["heading"]
            manuscript, attached_reference_list_from_content = manuscript_snapshot(outline_data, section_attached_references_content)
            update_job(
                message=f"Writing section {index} of {len(sections)}: {section_label}",
                current_section=section_label,
                completed_sections=index - 1,
                manuscript=manuscript,
                attached_reference_list_from_content=attached_reference_list_from_content,
            )
            agent = abstract_writer if section["is_abstract"] else writer
            section_content_pre_summary = content_pre_summary
            (
                raw_content,
                content_pre_summary,
                concept_map,
                section_attached_references_content,
                display_content,
            ) = await generateContent(
                agent,
                content_pre_summary,
                section["current_section"],
                section["instructions"],
                section_attached_references_content,
                attached_references_db,
            )
            raw_content = str(raw_content or "")
            content_pre_summary = str(content_pre_summary or "")
            display_content = display_content or raw_content
            display_content_overrides[tuple(section["path"])] = display_content
            _set_content_value(section["items"], OUTLINE_CONTENT_AI, raw_content)
            _set_content_value(
                section["items"],
                OUTLINE_CONTENT_PRE_SUMMARY,
                section_content_pre_summary if section["is_abstract"] else content_pre_summary,
            )
            if concept_map:
                _set_content_value(section["items"], OUTLINE_CONCEPT_MAP, concept_map)
            _write_outline_file(generated_file_id, outline_data)
            manuscript, attached_reference_list_from_content = manuscript_snapshot(outline_data, section_attached_references_content)
            update_job(
                completed_sections=index,
                manuscript=manuscript,
                attached_reference_list_from_content=attached_reference_list_from_content,
            )
            if pause_if_requested(outline_data, index):
                return

        _set_generated_file_status(generated_file_id, db.generated_files_status.SUCCESS.value)
        manuscript, attached_reference_list_from_content = manuscript_snapshot(outline_data, section_attached_references_content)
        update_job(
            status="completed",
            message="Content generation completed.",
            current_section="",
            worker_active=False,
            generated_file={**generated_file, "status": db.generated_files_status.SUCCESS.value},
            manuscript=manuscript,
            attached_reference_list_from_content=attached_reference_list_from_content,
        )
    except asyncio.CancelledError:
        if outline_data:
            manuscript, attached_reference_list_from_content = manuscript_snapshot(outline_data, job.get("attached_reference_list_from_content", []))
        else:
            manuscript, attached_reference_list_from_content = job.get("manuscript", []), job.get("attached_reference_list_from_content", [])
        _set_generated_file_status(generated_file_id, "cancelled")
        update_job(
            status="paused",
            message="Generation paused. Click Generate to continue with the remaining outline.",
            current_section="",
            worker_active=False,
            generated_file={**generated_file, "status": "cancelled"},
            manuscript=manuscript,
            attached_reference_list_from_content=attached_reference_list_from_content,
        )
    except Exception as exp:
        logger.exception("Unable to generate manuscript for generated file %s", generated_file_id)
        _set_generated_file_status(generated_file_id, "error")
        update_job(
            status="error",
            error=_friendly_error_detail("generate manuscript content", exp)[1],
            message="Content generation stopped.",
            current_section="",
            worker_active=False,
        )
    finally:
        job["task"] = None


async def _handle_health():
    checks = {
        "backend": _service_health("ok", "Backend", "Backend API is reachable."),
        "ai_model": _check_ai_model_health(),
        "chroma_db": _check_chroma_health(),
        "postgres": _check_postgres_health(),
    }
    statuses = {check["status"] for check in checks.values()}
    overall_status = "error" if "error" in statuses else "warning" if "warning" in statuses else "ok"
    return {
        "status": overall_status,
        "checks": checks,
        "default_model": Config.env_config.get("DEFAULT_AI_MODEL"),
        "chroma_host": Config.env_config.get("CHROMA_HOST"),
        "chroma_port": Config.env_config.get("CHROMA_PORT"),
        "checked_at": datetime.now().isoformat(),
    }


async def _handle_outline_templates():
    try:
        return {"templates": _outline_templates()}
    except Exception as exp:
        raise _api_error("load outline templates", exp)


async def _handle_outline_template(template_name):
    try:
        return {"name": template_name, "content": _outline_template_content(template_name)}
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("load outline template", exp)


async def _handle_uploaded_outline_templates(credentials_id):
    try:
        validated_credentials_id = _validated_credentials_id(credentials_id)
        return {"templates": _uploaded_outline_templates(validated_credentials_id)}
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("load uploaded outline templates", exp)


async def _handle_upload_outline_templates(files, credentials_id):
    try:
        validated_credentials_id = _validated_credentials_id(credentials_id)
        saved_templates = []
        seen_names: set[str] = set()
        for upload in files:
            file_name = _uploaded_outline_template_file_name(upload.filename)
            if file_name.casefold() in seen_names:
                continue
            seen_names.add(file_name.casefold())
            saved_templates.append(await _save_uploaded_outline_template(validated_credentials_id, upload))

        if not saved_templates:
            raise HTTPException(status_code=400, detail="Choose at least one Markdown or Word outline file to upload.")

        return {
            "message": "Outline templates uploaded successfully.",
            "templates": _uploaded_outline_templates(validated_credentials_id),
            "saved_templates": saved_templates,
        }
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("upload outline templates", exp)


async def _handle_uploaded_outline_template(template_name, credentials_id):
    try:
        validated_credentials_id = _validated_credentials_id(credentials_id)
        return {
            "name": template_name,
            "content": _uploaded_outline_template_content(validated_credentials_id, template_name),
        }
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("load uploaded outline template", exp)


async def _handle_update_uploaded_outline_template(template_name, request):
    try:
        validated_credentials_id = _validated_credentials_id(request.credentials_id)
        updated_template = _write_uploaded_outline_template_content(
            validated_credentials_id,
            template_name,
            request.outline,
        )
        return {
            "message": f"Outline template {updated_template['file_name']} saved.",
            "template": updated_template,
            "templates": _uploaded_outline_templates(validated_credentials_id),
        }
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("update uploaded outline template", exp)


async def _handle_rename_uploaded_outline_template(template_name, request):
    try:
        validated_credentials_id = _validated_credentials_id(request.credentials_id)
        renamed_template = _rename_uploaded_outline_template(
            validated_credentials_id,
            template_name,
            request.file_name,
        )
        return {
            "message": f"Outline template renamed.",
            "template": renamed_template,
            "templates": _uploaded_outline_templates(validated_credentials_id),
        }
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("rename uploaded outline template", exp)


async def _handle_delete_uploaded_outline_template(template_name, credentials_id):
    try:
        validated_credentials_id = _validated_credentials_id(credentials_id)
        deleted_template = _delete_uploaded_outline_template(validated_credentials_id, template_name)
        return {
            "message": f"Outline template {deleted_template['file_name']} removed.",
            "template": deleted_template,
            "templates": _uploaded_outline_templates(validated_credentials_id),
        }
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("delete uploaded outline template", exp)


async def _handle_workspace_data():
    return {
        "outline_template": '',
        "manuscript": [],
        "attached_reference_list_from_content": [],
        "generated_documents": [],
        "uploaded_documents": [],
    }


async def _handle_maintainer_logs(
    email,
    session,
    token,
    search,
    status,
    date_from,
    date_to,
    sort_by: LogSortField,
    sort_dir: SortDirection,
    page,
    page_size,
):
    try:
        validate_maintainer_access(email, session, token)
        entries = _log_entries()
        available_statuses = sorted({entry["status"] for entry in entries if entry.get("status")})
        normalized_search = (search or "").strip().lower()
        normalized_status = (status or "").strip().upper()
        from_date = _parse_log_date(date_from)
        to_date = _parse_log_date(date_to, end_of_day=True)

        filtered_entries = entries
        if normalized_search:
            filtered_entries = [
                entry
                for entry in filtered_entries
                if normalized_search in f"{entry.get('date', '')} {entry.get('status', '')} {entry.get('message', '')}".lower()
            ]
        if normalized_status and normalized_status != "ALL":
            filtered_entries = [entry for entry in filtered_entries if entry.get("status") == normalized_status]
        if from_date or to_date:
            next_entries = []
            for entry in filtered_entries:
                entry_date = _parse_log_date(entry.get("date"))
                if entry_date is None:
                    continue
                if from_date and entry_date < from_date:
                    continue
                if to_date and entry_date > to_date:
                    continue
                next_entries.append(entry)
            filtered_entries = next_entries

        reverse = sort_dir == "desc"
        filtered_entries = sorted(
            filtered_entries,
            key=lambda entry: str(entry.get(sort_by, "")).lower(),
            reverse=reverse,
        )

        total = len(filtered_entries)
        current_page = max(int(page), 1)
        current_page_size = max(min(int(page_size), 100), 10)
        total_pages = max((total + current_page_size - 1) // current_page_size, 1)
        if current_page > total_pages:
            current_page = total_pages
        start = (current_page - 1) * current_page_size
        end = start + current_page_size

        return {
            "entries": [_public_log_entry(entry) for entry in filtered_entries[start:end]],
            "total": total,
            "page": current_page,
            "page_size": current_page_size,
            "total_pages": total_pages,
            "statuses": available_statuses,
        }
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("load application logs", exp)


async def _handle_clear_maintainer_logs(request: MaintainerLogClearRequest, token):
    try:
        validate_maintainer_access(request.email, request.session, token)
        removed_count = _clear_log_entries(request.entry_ids)
        return {
            "message": f"Cleared {removed_count} log {'entry' if removed_count == 1 else 'entries'}.",
            "removed_count": removed_count,
        }
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("clear application logs", exp)


async def _handle_login(request):
    try:
        return authenticate_user(request.email, request.password)
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("authenticate user", exp)


async def _handle_azure_login(request):
    try:
        return auth_azure_login(request)
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("start Azure sign-in", exp)


async def _handle_azure_status():
    try:
        return auth_azure_status()
    except Exception as exp:
        raise _api_error("check Azure sign-in", exp)


async def _handle_azure_auth_callback(request, code, state, error):
    try:
        return auth_azure_callback(request, code=code, state=state, error=error)
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("complete Azure sign-in", exp)


async def _handle_azure_session(http_request, request):
    try:
        return authenticate_azure_session(http_request, request.code, request.state)
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("authenticate with Azure", exp)


async def _handle_create_account(request):
    try:
        return auth_create_account(request)
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("create account", exp)


async def _handle_forgot_password(request):
    try:
        return await send_password_reset_code(request.email)
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("send activation code", exp)


async def _handle_verify_reset_code(request):
    try:
        return auth_verify_reset_code(request.email, request.code)
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("verify activation code", exp)


async def _handle_reset_password(request):
    try:
        return auth_reset_password(request)
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("reset password", exp)


async def _handle_change_password(request):
    try:
        return auth_change_password(request)
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("change password", exp)


async def _handle_default_settings(session):
    try:
        return {**_create_guest_auth_payload(session=session), "llm_options": _llm_options()}
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("load default settings", exp)


async def _handle_get_settings(settings_id, email, session):
    try:
        from src import db

        owner_fields, owner_values, _, _ = _owner_filter_fields(email=email, session=session)
        df = db.selectFromDB(
            table_name="settings",
            field_names=["id", *owner_fields],
            field_values=[[settings_id], *owner_values],
            limit=1,
        )
        records = _records_from_dataframe(df)
        if not records:
            raise HTTPException(status_code=404, detail="Settings row was not found for this session.")

        return {"settings": _settings_record(records[0]), "llm_options": _llm_options()}
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("load settings", exp)


async def _handle_update_settings(settings_id, request, email, session):
    try:
        from src import db

        owner_fields, owner_values, _, _ = _owner_filter_fields(email=email, session=session)
        df = db.selectFromDB(
            table_name="settings",
            field_names=["id", *owner_fields],
            field_values=[[settings_id], *owner_values],
            limit=1,
        )
        if not _records_from_dataframe(df):
            raise HTTPException(status_code=404, detail="Settings row was not found for this session.")

        llm_options = _llm_options()
        valid_llms = {option["value"] for option in llm_options}
        if valid_llms and request.llm not in valid_llms:
            raise HTTPException(status_code=400, detail="Invalid LLM selection.")

        now = datetime.now()
        db.updateDB(
            table_name="settings",
            update_fields=["llm", "temperature", "instructions", "update_date"],
            update_values=[request.llm, request.temperature, request.instructions, now],
            select_fields=["id", *owner_fields],
            select_values=[[settings_id], *owner_values],
        )
        updated = db.selectFromDB(
            table_name="settings",
            field_names=["id", *owner_fields],
            field_values=[[settings_id], *owner_values],
            limit=1,
        )
        records = _records_from_dataframe(updated)
        return {"status": "saved", "settings": _settings_record(records[0]), "llm_options": llm_options}
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("update settings", exp)


async def _handle_generated_files(email, session, limit):
    try:
        records = _generated_records_for_owner(email=email, session=session, limit=limit)
        return {
            "generated_files": [_generated_file_record(record) for record in records],
            "generated_documents": [_generated_document_detail(record) for record in records],
        }
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("load generated files", exp)


async def _handle_create_generated_file(request):
    try:
        from src import db

        email = _validate_email(request.email) if request.email else ""
        session = "" if email else request.session.strip()
        if not email and not session:
            raise HTTPException(status_code=400, detail="A session is required to save a file.")
        file_name = request.file_name.strip()
        if not file_name:
            raise HTTPException(status_code=400, detail="File name is required.")

        if email:
            if request.settings_id is None:
                raise HTTPException(status_code=400, detail="Settings id is required for signed-in users.")
            owner_fields, owner_values, _, _ = _owner_filter_fields(email=email, session=session)
            settings = db.selectFromDB(
                table_name="settings",
                field_names=["id", *owner_fields],
                field_values=[[request.settings_id], *owner_values],
                limit=1,
            )
            if not _records_from_dataframe(settings):
                raise HTTPException(status_code=404, detail="Settings row was not found for this session.")

        valid_architectures = {item.value for item in db.generated_files_ai_architecture}
        if request.ai_architecture not in valid_architectures:
            raise HTTPException(status_code=400, detail="Invalid AI architecture.")

        now = datetime.now()
        replaced_generated_file = None
        duplicate = _duplicate_generated_file_for_owner(file_name, email=email, session=session)
        if duplicate:
            if not request.replace:
                raise HTTPException(
                    status_code=409,
                    detail={
                        "reason": "duplicate_generated_file",
                        "message": f"A generated document named {file_name} already exists. Replace it?",
                        "generated_file": _generated_file_record(duplicate),
                    },
                )
            _mark_generated_file_replaced(duplicate, now)
            replaced_generated_file = {
                **duplicate,
                "status": db.generated_files_status.DELETED.value,
                "update_date": now,
            }

        inserted_ids = db.insertIntoDB(
            table_name="generated_files",
            field_names=[
                "email",
                "session",
                "file_name",
                "status",
                "settings_id",
                "ai_architecture",
                "create_date",
                "update_date",
            ],
            field_values=[
                [email],
                [session],
                [file_name],
                [db.generated_files_status.CREATED.value],
                [request.settings_id],
                [request.ai_architecture],
                [now],
                [now],
            ],
        )
        generated_file = {
            "id": inserted_ids[0] if inserted_ids else None,
            "email": email,
            "session": session,
            "file_name": file_name,
            "status": db.generated_files_status.CREATED.value,
            "settings_id": request.settings_id,
            "ai_architecture": request.ai_architecture,
            "create_date": now,
            "update_date": now,
        }
        return {
            "status": "saved",
            "generated_file": _jsonable(generated_file),
            "replaced_generated_file": _jsonable(_generated_file_record(replaced_generated_file)) if replaced_generated_file else None,
        }
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("save generated file", exp)


async def _handle_update_generated_file(generated_file_id, request):
    try:
        from src import db

        file_name = request.file_name.strip()
        if not file_name:
            raise HTTPException(status_code=400, detail="File name is required.")

        existing = db.selectFromDB(
            table_name="generated_files",
            field_names=["id"],
            field_values=[[generated_file_id]],
            limit=1,
        )
        existing_records = _records_from_dataframe(existing)
        if not existing_records:
            raise HTTPException(status_code=404, detail="Generated file was not found.")
        existing_record = existing_records[0]

        now = datetime.now()
        replaced_generated_file = None
        duplicate = _duplicate_generated_file_for_owner(
            file_name,
            email=existing_record.get("email") or None,
            session=existing_record.get("session") or None,
            exclude_id=generated_file_id,
        )
        if duplicate:
            if not request.replace:
                raise HTTPException(
                    status_code=409,
                    detail={
                        "reason": "duplicate_generated_file",
                        "message": f"A generated document named {file_name} already exists. Replace it?",
                        "generated_file": _generated_file_record(duplicate),
                    },
                )
            _mark_generated_file_replaced(duplicate, now)
            replaced_generated_file = {
                **duplicate,
                "status": db.generated_files_status.DELETED.value,
                "update_date": now,
            }

        db.updateDB(
            table_name="generated_files",
            update_fields=["file_name", "update_date"],
            update_values=[file_name, now],
            select_fields=["id"],
            select_values=[[generated_file_id]],
        )
        updated = db.selectFromDB(
            table_name="generated_files",
            field_names=["id"],
            field_values=[[generated_file_id]],
            limit=1,
        )
        records = _records_from_dataframe(updated)
        return {
            "status": "saved",
            "generated_file": _generated_file_record(records[0]),
            "replaced_generated_file": _generated_file_record(replaced_generated_file) if replaced_generated_file else None,
        }
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("update generated file", exp)


async def _handle_delete_generated_file(generated_file_id, email, session):
    try:
        from src import db

        generated_file = _generated_file_by_id(generated_file_id, email=email, session=session)
        now = datetime.now()
        db.updateDB(
            table_name="generated_files",
            update_fields=["status", "update_date"],
            update_values=[db.generated_files_status.DELETED.value, now],
            select_fields=["id"],
            select_values=[[generated_file_id]],
        )

        active_collections = db.selectFromDB(
            table_name="vector_db_collections",
            field_names=["generated_files_id", "status"],
            field_values=[[generated_file_id], [db.vector_db_collections_status.ACTIVE.value]],
        )
        for collection in _records_from_dataframe(active_collections):
            _delete_vector_collection_record(collection, now)

        records = _generated_records_for_owner(
            email=generated_file.get("email") or email,
            session=generated_file.get("session") or session,
            limit=1000,
        )
        deleted_file_name = str(generated_file.get("file_name") or "Untitled")
        return {
            "status": "deleted",
            "generated_file": _generated_file_record(
                {
                    **generated_file,
                    "status": db.generated_files_status.DELETED.value,
                    "update_date": now,
                }
            ),
            "generated_documents": [_generated_document_detail(record) for record in records],
            "message": f'Generated document removed.',
        }
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("remove generated file", exp)


async def _handle_generated_file_manuscript(generated_file_id, email, session):
    try:
        record = _generated_file_by_id(generated_file_id, email=email, session=session)
        manuscript, source_exists = _manuscript_from_outline_file(generated_file_id)
        attached_reference_map = _attached_reference_map_for_generated_file(generated_file_id)
        manuscript, attached_reference_list_from_content = _process_manuscript_citations(manuscript, attached_reference_map)
        raw_outline, _ = _raw_outline_from_outline_file(generated_file_id)
        return {
            "generated_file": _generated_file_record(record),
            "manuscript": manuscript,
            "attached_reference_list_from_content": attached_reference_list_from_content,
            "literature_search": _active_literature_collection(generated_file_id),
            "uploaded_files_collection": _active_uploaded_files_collection(generated_file_id),
            "attached_files": _attached_uploaded_documents(generated_file_id),
            "outline": raw_outline,
            "query": _query_details_for_generated_file(generated_file_id),
            "source_exists": source_exists,
            "message": "" if source_exists else "No manuscript content has been saved for this file yet.",
        }
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("load manuscript", exp)


async def _handle_update_generated_file_paragraph(generated_file_id, request):
    try:
        from src import db

        generated_file = _generated_file_by_id(generated_file_id, email=request.email, session=request.session)
        outline_data = _read_outline_file(generated_file_id)
        node = _outline_node_for_path(outline_data, request.section_path) if request.section_path else None
        if node is None:
            node = _outline_node_for_heading_and_paragraph(
                outline_data,
                request.section_heading,
                request.raw_paragraph,
            )
        if node is None:
            raise HTTPException(status_code=404, detail="The selected section could not be found. Reload the manuscript and try again.")

        items = _content_items(node)
        if not items or not _section_needs_ai(items):
            raise HTTPException(status_code=400, detail="Only generated manuscript paragraphs can be updated.")

        current_raw_content = str(_content_value(items, OUTLINE_CONTENT_AI, "") or "")
        if not current_raw_content.strip():
            raise HTTPException(status_code=400, detail="This section does not have generated content to update yet.")

        replacement = ""
        if request.action != "Remove":
            _, _, ContentWriterArchitecture = _safe_architecture_classes()
            architecture_type = generated_file.get("ai_architecture") or db.generated_files_ai_architecture.BASE.value
            architecture_type, collection_name, collection_name_lit_search = _resolve_generation_architecture(
                generated_file_id=generated_file_id,
                architecture_type=architecture_type,
            )
            writer = ContentWriterArchitecture(
                model_name=_model_name(request.model_name),
                temperature=request.temperature,
                instructions=request.instructions,
                type=architecture_type,
                collection_name=collection_name,
                collection_name_lit_search=collection_name_lit_search,
            )
            section_body = _section_body_from_outline_node(node) or current_raw_content
            prompt = _paragraph_update_prompt(request.section_path, section_body, request.paragraph_index, request.action)
            response = await writer.ainvoke(
                _content_state(
                    ContentRequest(
                        current_section=prompt,
                        content_pre_summary="",
                        content_specific_instructions=_paragraph_update_instruction(request.action),
                        model_name=request.model_name,
                        temperature=request.temperature,
                        instructions=request.instructions,
                        architecture_type=architecture_type,
                        collection_name=collection_name,
                        collection_name_lit_search=collection_name_lit_search,
                    )
                )
            )
            replacement = _content_response_text(response)
            if not replacement:
                raise HTTPException(status_code=502, detail="The AI model did not return updated paragraph text. Please try again.")

        updated_raw_content = _replace_paragraph_in_section_body(
            current_raw_content,
            request.paragraph_index,
            request.raw_paragraph,
            replacement,
            request.action,
        )
        _set_content_value(items, OUTLINE_CONTENT_AI, updated_raw_content)
        _write_outline_file(generated_file_id, outline_data)
        db.updateDB(
            table_name="generated_files",
            update_fields=["update_date"],
            update_values=[datetime.now()],
            select_fields=["id"],
            select_values=[[generated_file_id]],
        )

        attached_reference_map = _attached_reference_map_for_generated_file(generated_file_id)
        manuscript, attached_reference_list_from_content = _display_manuscript_from_outline_tree(outline_data, attached_reference_map)
        raw_outline, _ = _raw_outline_from_outline_file(generated_file_id)
        action_message = {
            "Expand": "Paragraph expanded.",
            "Rephrase": "Paragraph rephrased.",
            "Remove": "Paragraph removed.",
        }[request.action]
        return {
            "generated_file": _generated_file_record({**generated_file, "update_date": datetime.now()}),
            "manuscript": manuscript,
            "attached_reference_list_from_content": attached_reference_list_from_content,
            "outline": raw_outline,
            "message": action_message,
        }
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("update manuscript paragraph", exp)


async def _handle_update_generated_file_section_content(generated_file_id, request):
    try:
        from src import db
        from src.common import rawCiteContent

        generated_file = _generated_file_by_id(generated_file_id, email=request.email, session=request.session)
        outline_data = _read_outline_file(generated_file_id)
        node = _outline_node_for_path(outline_data, request.section_path) if request.section_path else None
        if node is None:
            node = _outline_node_for_heading_and_paragraph(outline_data, request.section_heading, "")
        if node is None:
            raise HTTPException(status_code=404, detail="The selected section could not be found. Reload the manuscript and try again.")

        edited_section_content = str(request.section_content or "").strip()
        if not edited_section_content:
            raise HTTPException(status_code=400, detail="Section content cannot be empty.")

        items = _content_items(node)
        if not items:
            node[OUTLINE_CONTENT_KEY] = []
            items = node[OUTLINE_CONTENT_KEY]

        attached_references_db = _attached_reference_map_for_generated_file(generated_file_id)
        attached_reference_list_from_content = _normalize_attached_references_content(
            request.attached_reference_list_from_content
        )
        if not attached_reference_list_from_content:
            _current_manuscript, attached_reference_list_from_content = _display_manuscript_from_outline_tree(
                outline_data,
                attached_references_db,
            )

        try:
            raw_section_content = rawCiteContent(
                edited_section_content,
                attached_reference_list_from_content,
                attached_references_db,
            )
        except Exception as exp:
            logger.exception("Unable to convert formatted citations while saving edited section.")
            raise HTTPException(
                status_code=400,
                detail="The edited section contains a citation that could not be matched to the current references. Check the citation numbers and try again.",
            ) from exp

        content_type = OUTLINE_CONTENT_AI
        if not _has_content_value(items, OUTLINE_CONTENT_AI) and _has_content_value(items, OUTLINE_CONTENT_USER):
            content_type = OUTLINE_CONTENT_USER
        _set_content_value(items, content_type, raw_section_content)
        _write_outline_file(generated_file_id, outline_data)

        now = datetime.now()
        db.updateDB(
            table_name="generated_files",
            update_fields=["update_date"],
            update_values=[now],
            select_fields=["id"],
            select_values=[[generated_file_id]],
        )

        manuscript, attached_reference_list_from_content = _display_manuscript_from_outline_tree(outline_data, attached_references_db)
        raw_outline, _ = _raw_outline_from_outline_file(generated_file_id)
        return {
            "generated_file": _generated_file_record({**generated_file, "update_date": now}),
            "manuscript": manuscript,
            "attached_reference_list_from_content": attached_reference_list_from_content,
            "outline": raw_outline,
            "message": "Section changes saved.",
        }
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("save manuscript section", exp)


async def _handle_download_generated_file(generated_file_id, download_format, email, session):
    try:
        from src import common as common_module

        record = _generated_file_by_id(generated_file_id, email=email, session=session)
        if not _outline_file_path(generated_file_id).exists():
            raise HTTPException(status_code=404, detail="No manuscript content has been saved for this file yet.")

        uploaded_collection = _active_uploaded_files_collection_record(generated_file_id)
        literature_collection = _active_literature_collection_record(generated_file_id)
        uploaded_collection_id = int(uploaded_collection["id"]) if uploaded_collection else None
        literature_collection_id = int(literature_collection["id"]) if literature_collection else None

        def build_doc_content() -> tuple[str, Any, str, str]:
            def get_attached_refs_for_download(
                vector_db_collections_id_uploaded_files: int | None,
                vector_db_collections_id_literature: int | None,
            ) -> tuple[list[Any], dict[str, Any]]:
                files: list[Any] = []
                file_info: dict[str, Any] = {}
                for collection_id in (vector_db_collections_id_uploaded_files, vector_db_collections_id_literature):
                    refs, info = common_module.getVectorDBFiles(collection_id)
                    files.extend(refs)
                    file_info.update(info)
                return files, file_info

            with DOC_CONTENT_LOCK:
                original_get_attached_refs = getattr(common_module, "getAttachedRefs", None)
                common_module.getAttachedRefs = get_attached_refs_for_download
                try:
                    return common_module.getDocContent(
                        generated_file_id,
                        uploaded_collection_id,
                        literature_collection_id,
                    )
                finally:
                    if original_get_attached_refs is not None:
                        common_module.getAttachedRefs = original_get_attached_refs

        content_md, content_docx, content_tex, bibs = await asyncio.to_thread(build_doc_content)

        file_stem = _safe_download_stem(str(record.get("file_name") or "generated-document"))
        file_name = f"{file_stem}.{download_format if download_format != 'latex' else 'zip'}"
        headers = {"Content-Disposition": f'attachment; filename="{file_name}"'}

        if download_format == "docx":
            output = BytesIO()
            content_docx.save(output)
            output.seek(0)
            return StreamingResponse(
                output,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers=headers,
            )

        if download_format == "latex":
            output = BytesIO()
            with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
                archive.writestr("main.tex", content_tex)
                archive.writestr("bibliography.bib", bibs or "")
            output.seek(0)
            return Response(content=output.getvalue(), media_type="application/zip", headers=headers)

        text_by_format = {
            "md": (content_md, "text/markdown; charset=utf-8"),
        }
        content, media_type = text_by_format[download_format]
        return Response(content=content, media_type=media_type, headers=headers)
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("download generated document", exp)


async def _handle_enable_generated_file_literature_search(generated_file_id, request):
    try:
        from src import db

        generated_file = _generated_file_by_id(generated_file_id, email=request.email, session=request.session)
        existing_literature_collection = _active_literature_collection_record(generated_file_id)
        if existing_literature_collection:
            _delete_vector_collection_record(existing_literature_collection)

        literature_collection = _create_literature_collection(generated_file)
        collection_name = literature_collection["collection_name"]

        now = datetime.now()
        manuscript, concept_maps, attached_reference_list_from_content, source_exists = _reset_generated_document_content(generated_file_id)
        db.updateDB(
            table_name="generated_files",
            update_fields=["ai_architecture", "status", "update_date"],
            update_values=[db.generated_files_ai_architecture.RAG.value, db.generated_files_status.CREATED.value, now],
            select_fields=["id"],
            select_values=[[generated_file_id]],
        )
        updated_file = {
            **generated_file,
            "ai_architecture": db.generated_files_ai_architecture.RAG.value,
            "status": db.generated_files_status.CREATED.value,
            "update_date": now,
        }
        return {
            "status": "enabled",
            "generated_file": _jsonable(updated_file),
            "literature_search": _jsonable(literature_collection),
            "collection_name": collection_name,
            "manuscript": manuscript,
            "concept_maps": concept_maps,
            "attached_reference_list_from_content": attached_reference_list_from_content,
            "content_reset": source_exists,
            "message": (
                "Literature Search enabled. The manuscript content was reset because the generation context changed."
                if source_exists
                else "Literature Search is enabled for this file."
            ),
        }
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("enable literature search", exp)


async def _handle_disable_generated_file_literature_search(generated_file_id, email, session):
    try:
        from src import db

        generated_file = _generated_file_by_id(generated_file_id, email=email, session=session)
        literature_collection = _active_literature_collection_record(generated_file_id)
        if not literature_collection:
            uploaded_collection = _active_uploaded_files_collection_record(generated_file_id)
            has_uploaded_collection = bool(
                uploaded_collection and _attached_uploaded_files(int(uploaded_collection["id"]))
            )
            next_architecture = (
                db.generated_files_ai_architecture.RAG.value
                if has_uploaded_collection
                else db.generated_files_ai_architecture.BASE.value
            )
            now = datetime.now()
            if generated_file.get("ai_architecture") != next_architecture:
                db.updateDB(
                    table_name="generated_files",
                    update_fields=["ai_architecture", "update_date"],
                    update_values=[next_architecture, now],
                    select_fields=["id"],
                    select_values=[[generated_file_id]],
                )
            updated_file = {
                **generated_file,
                "ai_architecture": next_architecture,
                "update_date": now,
            }
            return {
                "status": "disabled",
                "generated_file": _generated_file_record(updated_file),
                "literature_search": None,
                "uploaded_files_collection": _active_uploaded_files_collection(generated_file_id),
                "attached_files": _attached_uploaded_documents(generated_file_id),
                "message": "Literature Search is already disabled for this file.",
            }

        now = datetime.now()
        _delete_vector_collection_record(literature_collection, now)

        uploaded_collection = _active_uploaded_files_collection_record(generated_file_id)
        has_uploaded_collection = bool(
            uploaded_collection and _attached_uploaded_files(int(uploaded_collection["id"]))
        )
        next_architecture = (
            db.generated_files_ai_architecture.RAG.value
            if has_uploaded_collection
            else db.generated_files_ai_architecture.BASE.value
        )
        manuscript, concept_maps, attached_reference_list_from_content, source_exists = _reset_generated_document_content(generated_file_id)
        db.updateDB(
            table_name="generated_files",
            update_fields=["ai_architecture", "status", "update_date"],
            update_values=[next_architecture, db.generated_files_status.CREATED.value, now],
            select_fields=["id"],
            select_values=[[generated_file_id]],
        )

        updated_file = {
            **generated_file,
            "ai_architecture": next_architecture,
            "status": db.generated_files_status.CREATED.value,
            "update_date": now,
        }
        return {
            "status": "disabled",
            "generated_file": _generated_file_record(updated_file),
            "literature_search": None,
            "uploaded_files_collection": _active_uploaded_files_collection(generated_file_id),
            "attached_files": _attached_uploaded_documents(generated_file_id),
            "manuscript": manuscript,
            "concept_maps": concept_maps,
            "attached_reference_list_from_content": attached_reference_list_from_content,
            "content_reset": source_exists,
            "message": (
                "Literature Search disabled. The manuscript content was reset because the generation context changed."
                if source_exists
                else "Literature Search disabled for this file."
            ),
        }
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("disable literature search", exp)


async def _handle_attach_uploaded_files_to_generated_file(generated_file_id, request):
    try:
        from src import db

        uploaded_file_ids = sorted({int(file_id) for file_id in request.uploaded_file_ids})
        if not uploaded_file_ids:
            raise HTTPException(status_code=400, detail="Select at least one uploaded document to attach.")

        generated_file = _generated_file_by_id(generated_file_id, email=request.email, session=request.session)
        uploaded_records = _uploaded_file_records_by_ids(
            uploaded_file_ids,
            email=generated_file.get("email") or request.email,
            session=generated_file.get("session") or request.session,
        )
        found_ids = {int(record["id"]) for record in uploaded_records}
        missing_ids = [file_id for file_id in uploaded_file_ids if file_id not in found_ids]
        if missing_ids:
            raise HTTPException(status_code=404, detail="One or more selected uploaded documents were not found.")

        active_collection = _active_uploaded_files_collection_record(generated_file_id)
        if active_collection:
            attached_records = _attached_uploaded_files(int(active_collection["id"]))
            attached_ids = {int(record["id"]) for record in attached_records}
            duplicate_records = [record for record in uploaded_records if int(record["id"]) in attached_ids]
            if duplicate_records:
                duplicate_names = [str(record.get("file_name") or "Untitled") for record in duplicate_records]
                raise HTTPException(
                    status_code=409,
                    detail={
                        "reason": "already_attached",
                        "message": f"Already attached: {', '.join(duplicate_names)}.",
                        "duplicates": duplicate_names,
                    },
                )

            if attached_records and request.mode == "ask":
                attached_names = [str(record.get("file_name") or "Untitled") for record in attached_records]
                selected_names = [str(record.get("file_name") or "Untitled") for record in uploaded_records]
                raise HTTPException(
                    status_code=409,
                    detail={
                        "reason": "existing_attachments",
                        "message": "This generated document already has attached files.",
                        "attached_files": attached_names,
                        "selected_files": selected_names,
                    },
                )

            if attached_records and request.mode == "replace":
                now = datetime.now()
                _delete_vector_collection_record(active_collection, now)
                active_collection = _create_uploaded_files_collection(generated_file)
        else:
            active_collection = _create_uploaded_files_collection(generated_file)

        vector_db_collections_id = int(active_collection["id"])
        collection_name = active_collection.get("collection_name") or _vector_collection_name(vector_db_collections_id)
        file_paths = _uploaded_file_paths(uploaded_records)
        _insert_uploaded_file_attachment_rows(vector_db_collections_id, uploaded_records)
        _load_uploaded_files_to_vector_collection(collection_name, file_paths)

        now = datetime.now()
        db.updateDB(
            table_name="generated_files",
            update_fields=["ai_architecture", "update_date"],
            update_values=[db.generated_files_ai_architecture.RAG.value, now],
            select_fields=["id"],
            select_values=[[generated_file_id]],
        )
        db.updateDB(
            table_name="vector_db_collections",
            update_fields=["update_date"],
            update_values=[now],
            select_fields=["id"],
            select_values=[[vector_db_collections_id]],
        )

        updated_file = {
            **generated_file,
            "ai_architecture": db.generated_files_ai_architecture.RAG.value,
            "update_date": now,
        }
        attached_records = _attached_uploaded_files(vector_db_collections_id)
        manuscript, concept_maps, attached_reference_list_from_content, source_exists = _reset_generated_document_content(generated_file_id)
        return {
            "status": "attached",
            "generated_file": _jsonable(updated_file),
            "collection": _jsonable(
                {
                    **active_collection,
                    "collection_name": collection_name,
                    "update_date": now,
                }
            ),
            "attached_files": [_uploaded_document(record) for record in attached_records],
            "manuscript": manuscript,
            "concept_maps": concept_maps,
            "attached_reference_list_from_content": attached_reference_list_from_content,
            "content_reset": source_exists,
            "message": (
                "Uploaded documents attached. The manuscript content was reset because the attached references changed."
                if source_exists
                else "Uploaded documents attached to this generated document."
            ),
        }
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("attach uploaded documents", exp)


async def _handle_remove_uploaded_file_attachment(generated_file_id, uploaded_file_id, email, session):
    try:
        from src import db

        generated_file = _generated_file_by_id(generated_file_id, email=email, session=session)
        active_collection = _active_uploaded_files_collection_record(generated_file_id)
        if not active_collection:
            raise HTTPException(status_code=404, detail="There are no uploaded files attached to this generated document.")

        attached_records = _attached_uploaded_files(int(active_collection["id"]))
        attached_ids = {int(record["id"]) for record in attached_records}
        if int(uploaded_file_id) not in attached_ids:
            raise HTTPException(status_code=404, detail="That uploaded file is not attached to this generated document.")

        remaining_records = [record for record in attached_records if int(record["id"]) != int(uploaded_file_id)]
        now = datetime.now()
        _delete_vector_collection_record(active_collection, now)

        uploaded_files_collection = None
        if remaining_records:
            uploaded_files_collection = _create_uploaded_files_collection_from_records(generated_file, remaining_records)

        has_literature_collection = _active_literature_collection_record(generated_file_id) is not None
        next_architecture = (
            db.generated_files_ai_architecture.RAG.value
            if remaining_records or has_literature_collection
            else db.generated_files_ai_architecture.BASE.value
        )
        manuscript, concept_maps, attached_reference_list_from_content, source_exists = _reset_generated_document_content(generated_file_id)
        db.updateDB(
            table_name="generated_files",
            update_fields=["ai_architecture", "status", "update_date"],
            update_values=[next_architecture, db.generated_files_status.CREATED.value, now],
            select_fields=["id"],
            select_values=[[generated_file_id]],
        )

        updated_file = {
            **generated_file,
            "ai_architecture": next_architecture,
            "status": db.generated_files_status.CREATED.value,
            "update_date": now,
        }
        return {
            "status": "removed",
            "generated_file": _jsonable(updated_file),
            "uploaded_files_collection": _jsonable(uploaded_files_collection),
            "attached_files": [_uploaded_document(record) for record in remaining_records],
            "manuscript": manuscript,
            "concept_maps": concept_maps,
            "attached_reference_list_from_content": attached_reference_list_from_content,
            "content_reset": source_exists,
            "message": (
                "Uploaded document removed. The manuscript content was reset because the attached references changed."
                if source_exists
                else "Uploaded document removed from this generated document."
            ),
        }
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("remove attached uploaded document", exp)


async def _handle_generated_file_concept_map(generated_file_id, email, session):
    try:
        record = _generated_file_by_id(generated_file_id, email=email, session=session)
        concept_maps, source_exists = _concept_maps_from_outline_file(generated_file_id)
        return {
            "generated_file": _generated_file_record(record),
            "concept_maps": concept_maps,
            "source_exists": source_exists,
            "message": "" if concept_maps else "No concept map has been generated for this file yet.",
        }
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("load concept map", exp)


async def _handle_generate_generated_file(generated_file_id, request):
    try:
        active_job = _active_generation_job_for_file(generated_file_id)
        if active_job:
            detail = (
                "Generation is pausing. Please try again in a moment."
                if active_job.get("pause_requested")
                else "Content generation is already running for this file."
            )
            raise HTTPException(status_code=409, detail=detail)

        _required_default_model()
        generated_file, processed_outline, outline_file_path = _save_processed_outline(
            generated_file_id,
            request.outline,
            email=request.email,
            session=request.session,
            preserve_generated_content=request.mode == "remaining",
        )
        if request.mode == "restart":
            _refresh_vector_collections_for_regeneration(generated_file)

        sections = _extract_generation_sections(processed_outline, remaining_only=request.mode == "remaining")
        if not sections:
            from src import db

            _set_generated_file_status(generated_file_id, db.generated_files_status.SUCCESS.value)
            completed_file = {
                **generated_file,
                "status": db.generated_files_status.SUCCESS.value,
            }
            manuscript, attached_reference_list_from_content = _display_manuscript_from_outline_tree(
                processed_outline,
                _attached_reference_map_for_generated_file(generated_file_id),
            )
            completed_at = datetime.now().isoformat()
            job = {
                "id": "",
                "generated_file_id": generated_file_id,
                "status": "completed",
                "message": "Content generation completed.",
                "error": "",
                "current_section": "",
                "completed_sections": 0,
                "total_sections": 0,
                "mode": request.mode,
                "manuscript": manuscript,
                "attached_reference_list_from_content": attached_reference_list_from_content,
                "generated_file": _generated_file_record(completed_file),
                "outline_path": str(outline_file_path),
                "created_at": completed_at,
                "updated_at": completed_at,
            }
            return {"status": "completed", "job": _job_snapshot(job)}

        job_id = uuid4().hex
        manuscript = _manuscript_from_outline_tree(processed_outline)
        job = {
            "id": job_id,
            "generated_file_id": generated_file_id,
            "status": "queued",
            "message": "Structured outline saved. Content generation is queued...",
            "error": "",
            "current_section": "",
            "completed_sections": 0,
            "total_sections": 0,
            "mode": request.mode,
            "manuscript": manuscript,
            "attached_reference_list_from_content": [],
            "generated_file": _generated_file_record(generated_file),
            "outline_path": str(outline_file_path),
            "pause_requested": False,
            "worker_active": True,
            "task": None,
            "created_at": datetime.now().isoformat(),
            "updated_at": datetime.now().isoformat(),
        }
        GENERATION_JOBS[job_id] = job
        task = asyncio.create_task(_run_generation_job(job_id, generated_file_id, request, _generated_file_record(generated_file)))
        task.add_done_callback(lambda _task, current_job=job: _finalize_generation_task(current_job))
        job["task"] = task
        return {"status": "queued", "job": _job_snapshot(job)}
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("start content generation", exp)


async def _handle_generation_job(job_id):
    job = GENERATION_JOBS.get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="That generation job was not found. Start generation again to create a new job.")
    return {"job": _job_snapshot(job)}


async def _handle_pause_generation_job(job_id):
    job = GENERATION_JOBS.get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="That generation job was not found.")

    if job.get("status") in {"queued", "running"}:
        job["pause_requested"] = True
        job["status"] = "paused"
        job["message"] = "Generation paused. Click Generate to continue with the remaining outline."
        job["current_section"] = ""
        job["updated_at"] = datetime.now().isoformat()
        if job.get("generated_file_id"):
            _set_generated_file_status(int(job["generated_file_id"]), "cancelled")
        if isinstance(job.get("generated_file"), dict):
            job["generated_file"] = {**job["generated_file"], "status": "cancelled"}
        task = job.get("task")
        if isinstance(task, asyncio.Task) and not task.done():
            task.cancel()
        else:
            _finalize_generation_task(job)

    return {"job": _job_snapshot(job)}


async def _handle_uploaded_files(email, session, limit):
    try:
        records = _uploaded_records_for_owner(email=email, session=session, limit=limit)
        return {
            "uploaded_files": [_uploaded_file_record(record) for record in records],
            "uploaded_documents": [_uploaded_document(record) for record in records],
        }
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("load uploaded files", exp)


async def _handle_upload_uploaded_files(files, email, session, replace):
    try:
        from src import db

        normalized_email, normalized_session, _, _ = _owner_identity(email=email, session=session)
        incoming_files: list[tuple[UploadFile, str]] = []
        seen_names: set[str] = set()
        for upload in files:
            file_name = _safe_uploaded_file_name(upload.filename)
            name_key = file_name.casefold()
            if name_key in seen_names:
                continue
            seen_names.add(name_key)
            incoming_files.append((upload, file_name))

        if not incoming_files:
            raise HTTPException(status_code=400, detail="Choose at least one document to upload.")

        existing_by_name = {
            file_name: existing
            for _, file_name in incoming_files
            if (existing := _uploaded_file_by_owner_and_name(file_name, email=normalized_email, session=normalized_session))
        }
        if existing_by_name and not replace:
            duplicate_names = sorted(existing_by_name)
            raise HTTPException(
                status_code=409,
                detail={
                    "message": (
                        "One or more uploaded documents already exist. Confirm replacement to update the saved files."
                    ),
                    "duplicates": duplicate_names,
                },
            )

        now = datetime.now()
        uploaded_records: list[dict[str, Any]] = []
        for upload, file_name in incoming_files:
            existing = existing_by_name.get(file_name)
            if existing:
                uploaded_file_id = int(existing["id"])
                db.updateDB(
                    table_name="uploaded_files",
                    update_fields=["status", "update_date"],
                    update_values=[db.uploaded_files_status.UPLOADED.value, now],
                    select_fields=["id"],
                    select_values=[[uploaded_file_id]],
                )
                record = {
                    **existing,
                    "status": db.uploaded_files_status.UPLOADED.value,
                    "update_date": now,
                }
            else:
                inserted_ids = db.insertIntoDB(
                    table_name="uploaded_files",
                    field_names=["email", "session", "file_name", "status", "create_date", "update_date"],
                    field_values=[
                        [normalized_email],
                        [normalized_session],
                        [file_name],
                        [db.uploaded_files_status.UPLOADED.value],
                        [now],
                        [now],
                    ],
                )
                uploaded_file_id = int(inserted_ids[0])
                record = {
                    "id": uploaded_file_id,
                    "email": normalized_email,
                    "session": normalized_session,
                    "file_name": file_name,
                    "status": db.uploaded_files_status.UPLOADED.value,
                    "create_date": now,
                    "update_date": now,
                }

            saved_path = _save_uploaded_file(upload, uploaded_file_id, file_name)
            uploaded_records.append({**record, "path": str(saved_path)})

        records = _uploaded_records_for_owner(email=normalized_email, session=normalized_session, limit=200)
        return {
            "message": "Documents uploaded successfully.",
            "uploaded_files": [_uploaded_file_record(record) for record in records],
            "uploaded_documents": [_uploaded_document(record) for record in records],
            "saved_files": [_jsonable(record) for record in uploaded_records],
        }
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("upload documents", exp)


async def _handle_update_uploaded_file(uploaded_file_id, request):
    try:
        from src import db

        normalized_email, normalized_session, _, _ = _owner_identity(email=request.email, session=request.session)
        file_name = _safe_uploaded_file_name(request.file_name)
        uploaded_records = _uploaded_file_records_by_ids(
            [uploaded_file_id],
            email=normalized_email,
            session=normalized_session,
        )
        if not uploaded_records:
            raise HTTPException(status_code=404, detail="Uploaded document was not found.")

        uploaded_record = uploaded_records[0]
        duplicate = _uploaded_file_by_owner_and_name(
            file_name,
            email=normalized_email,
            session=normalized_session,
        )
        if duplicate and int(duplicate["id"]) != int(uploaded_file_id):
            raise HTTPException(status_code=409, detail=f"An uploaded document named {file_name} already exists.")

        now = datetime.now()
        db.updateDB(
            table_name="uploaded_files",
            update_fields=["file_name", "update_date"],
            update_values=[file_name, now],
            select_fields=["id"],
            select_values=[[uploaded_file_id]],
        )

        records = _uploaded_records_for_owner(email=normalized_email, session=normalized_session, limit=200)
        updated_record = {
            **uploaded_record,
            "file_name": file_name,
            "update_date": now,
        }
        return {
            "status": "saved",
            "uploaded_file": _uploaded_file_record(updated_record),
            "uploaded_document": _uploaded_document(updated_record),
            "uploaded_documents": [_uploaded_document(record) for record in records],
            "affected_documents": [],
            "message": f"Uploaded document renamed.",
        }
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("rename uploaded document", exp)


async def _handle_delete_uploaded_file(uploaded_file_id, email, session):
    try:
        from src import db

        normalized_email, normalized_session, _, _ = _owner_identity(email=email, session=session)
        uploaded_records = _uploaded_file_records_by_ids(
            [uploaded_file_id],
            email=normalized_email,
            session=normalized_session,
        )
        if not uploaded_records:
            raise HTTPException(status_code=404, detail="Uploaded document was not found.")

        uploaded_record = uploaded_records[0]
        affected_collections = _active_uploaded_file_collections_for_uploaded_file(uploaded_file_id)
        affected_documents: list[dict[str, Any]] = []
        now = datetime.now()

        for collection in affected_collections:
            generated_file_id = int(collection["generated_files_id"])
            generated_file = _generated_file_by_id(
                generated_file_id,
                email=normalized_email or None,
                session=normalized_session or None,
            )
            attached_records = _attached_uploaded_files(int(collection["id"]))
            remaining_records = [
                record
                for record in attached_records
                if int(record["id"]) != int(uploaded_file_id)
                and record.get("status") != db.uploaded_files_status.DELETED.value
            ]

            _delete_vector_collection_record(collection, now)

            uploaded_files_collection = None
            if remaining_records:
                uploaded_files_collection = _create_uploaded_files_collection_from_records(generated_file, remaining_records)

            has_literature_collection = _active_literature_collection_record(generated_file_id) is not None
            next_architecture = (
                db.generated_files_ai_architecture.RAG.value
                if remaining_records or has_literature_collection
                else db.generated_files_ai_architecture.BASE.value
            )
            manuscript, concept_maps, attached_reference_list_from_content, source_exists = _reset_generated_document_content(generated_file_id)
            db.updateDB(
                table_name="generated_files",
                update_fields=["ai_architecture", "status", "update_date"],
                update_values=[next_architecture, db.generated_files_status.CREATED.value, now],
                select_fields=["id"],
                select_values=[[generated_file_id]],
            )

            updated_file = {
                **generated_file,
                "ai_architecture": next_architecture,
                "status": db.generated_files_status.CREATED.value,
                "update_date": now,
            }
            affected_documents.append(
                {
                    "generated_file": _generated_file_record(updated_file),
                    "uploaded_files_collection": _jsonable(uploaded_files_collection),
                    "attached_files": [_uploaded_document(record) for record in remaining_records],
                    "manuscript": manuscript,
                    "concept_maps": concept_maps,
                    "attached_reference_list_from_content": attached_reference_list_from_content,
                    "content_reset": source_exists,
                }
            )

        db.updateDB(
            table_name="uploaded_files",
            update_fields=["status", "update_date"],
            update_values=[db.uploaded_files_status.DELETED.value, now],
            select_fields=["id"],
            select_values=[[uploaded_file_id]],
        )

        uploaded_path = _uploaded_doc_path_by_id(int(uploaded_file_id), str(uploaded_record.get("file_name") or ""))
        try:
            uploaded_path.unlink(missing_ok=True)
        except Exception:
            logger.warning("Unable to remove uploaded file from disk: %s", uploaded_path, exc_info=True)

        records = _uploaded_records_for_owner(email=normalized_email, session=normalized_session, limit=200)
        affected_count = len(affected_documents)
        return {
            "status": "deleted",
            "uploaded_file": _uploaded_file_record(
                {
                    **uploaded_record,
                    "status": db.uploaded_files_status.DELETED.value,
                    "update_date": now,
                }
            ),
            "uploaded_documents": [_uploaded_document(record) for record in records],
            "affected_documents": affected_documents,
            "message": (
                f"Uploaded document deleted. {affected_count} generated document"
                f"{'' if affected_count == 1 else 's'} reset because the attached references changed."
                if affected_count
                else "Uploaded document deleted."
            ),
        }
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("delete uploaded document", exp)


async def _handle_create_outline(request):
    try:
        from src.ai.architecture import OutlineCreatorArchitecture
        from src.generate import generateOutline

        payload = await _outline_payload_from_request(request)
        if not payload["query"]:
            raise HTTPException(status_code=400, detail="Write a query before creating an outline.")

        _save_query_text(payload["generated_file_id"], payload["query"])
        details = await _outline_reference_documents_details(payload["reference_documents"], payload["generated_file_id"])
        architecture = OutlineCreatorArchitecture(
            model_name=_model_name(payload["model_name"]),
            temperature=payload["temperature"],
            instructions=payload["instructions"],
        )
        content = await generateOutline(architecture, query=payload["query"], details=details)
        return {"result": {"content": _jsonable(content)}}
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("create outline", exp)


async def _handle_import_outline(file, credentials_id):
    try:
        validated_credentials_id = _validated_credentials_id(credentials_id)
        if file is None or not file.filename:
            raise HTTPException(status_code=400, detail="Choose a Markdown or Word outline file to upload.")

        saved_template = await _save_uploaded_outline_template(validated_credentials_id, file)
        return {
            "message": f"Outline imported from {file.filename} and saved as {saved_template['file_name']}.",
            "result": {
                "content": saved_template["content"],
                "file_name": saved_template["file_name"],
                "template": saved_template,
            },
            "templates": _uploaded_outline_templates(validated_credentials_id),
        }
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("import outline", exp)


async def _handle_format_outline(request):
    try:
        from src.ai.architecture import OutlineFormatterArchitecture

        architecture = OutlineFormatterArchitecture(
            model_name=_model_name(request.model_name),
            temperature=request.temperature,
            instructions=request.instructions,
        )
        response = await architecture.ainvoke(
            _outline_state(query=request.query, outline_unstructured=request.outline_unstructured)
        )
        return {"result": _jsonable(response)}
    except HTTPException:
        raise
    except Exception as exp:
        raise _api_error("format outline", exp)


__all__ = [name for name in globals() if name.startswith("_") or name.isupper()]
